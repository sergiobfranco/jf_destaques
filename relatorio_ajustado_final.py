import os
import re
import shutil
from annotated_types import doc
from docx import Document
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google.oauth2 import service_account
from datetime import datetime
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def converter_asteriscos_em_negrito(paragraph):
    """
    Converte texto entre asteriscos (*texto*) em negrito no Word,
    mantendo os asteriscos visíveis
    """
    texto_completo = paragraph.text
    
    # Regex para encontrar texto entre asteriscos
    # Captura *texto* mas não **
    padrao_asterisco = r'\*([^\*]+)\*'
    
    # Limpar o parágrafo atual
    _limpar_paragrafo(paragraph)
    
    # Inicializar posição
    ultima_posicao = 0
    
    # Encontrar todas as ocorrências
    for match in re.finditer(padrao_asterisco, texto_completo):
        inicio, fim = match.span()
        texto_entre_asteriscos = match.group(1)  # Texto sem os asteriscos
        
        # Adicionar texto normal antes do próximo asterisco
        if inicio > ultima_posicao:
            paragraph.add_run(texto_completo[ultima_posicao:inicio])
        
        # Adicionar asterisco inicial
        paragraph.add_run("*")
        
        # Adicionar texto em negrito
        run_negrito = paragraph.add_run(texto_entre_asteriscos)
        run_negrito.bold = True
        
        # Forçar negrito no XML
        b = OxmlElement('w:b')
        b.set(qn('w:val'), 'on')
        run_negrito._element.get_or_add_rPr().append(b)
        
        # Adicionar asterisco final
        paragraph.add_run("*")
        
        # Atualizar posição
        ultima_posicao = fim
        
        print(f"   DEBUG: Aplicando negrito em: '*{texto_entre_asteriscos}*'")
    
    # Adicionar texto restante após o último asterisco
    if ultima_posicao < len(texto_completo):
        paragraph.add_run(texto_completo[ultima_posicao:])
    
    return True

    
    # Adicionar texto restante após o último asterisco (se houver)
    if ultima_posicao < len(texto_completo):
        paragraph.add_run(texto_completo[ultima_posicao:])
    
    return True

def _limpar_paragrafo(paragraph):
    # remove todos os runs atuais
    for r in paragraph.runs[::-1]:
        paragraph._p.remove(r._element)


def _add_text_preservando_asteriscos(paragraph, texto):
    """
    Adiciona texto ao parágrafo criando runs. Se encontrar trechos entre asteriscos
    (*texto*), mantém os asteriscos e aplica negrito ao texto interno.
    Não limpa o parágrafo; apenas acrescenta runs.
    """
    if not texto:
        return

    padrao = re.compile(r'\*([^\*]+)\*')
    ultima = 0
    for m in padrao.finditer(texto):
        inicio, fim = m.span()
        # texto antes
        if inicio > ultima:
            paragraph.add_run(texto[ultima:inicio])

        # asterisco inicial
        paragraph.add_run('*')

        # texto em negrito
        inner = m.group(1)
        run_b = paragraph.add_run(inner)
        run_b.bold = True

        # asterisco final
        paragraph.add_run('*')

        ultima = fim

    # restante
    if ultima < len(texto):
        paragraph.add_run(texto[ultima:])

# Funcao para upload no Google Drive
def upload_para_google_drive(caminho_arquivo, nome_arquivo, pasta_id):
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload

    SERVICE_ACCOUNT_FILE = 'service_account.json'
    SCOPES = ['https://www.googleapis.com/auth/drive']

    creds = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES
    )

    service = build('drive', 'v3', credentials=creds)

    file_metadata = {
        'name': nome_arquivo,
        'parents': [pasta_id]
    }

    media = MediaFileUpload(caminho_arquivo, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    uploaded_file = service.files().create(
        body=file_metadata,
        media_body=media,
        fields='id',
        supportsAllDrives=True  # ESSENCIAL para pastas compartilhadas
    ).execute()

    print(f"Upload concluido com ID: {uploaded_file.get('id')}")

def adicionar_hyperlink(paragraph, url, texto_display):
    """
    Adiciona um hyperlink a um paragrafo no documento Word
    """
    # Criar o elemento hyperlink
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), paragraph.part.relate_to(url,
                                                       "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                                                       is_external=True))

    # Criar o run com o texto do link
    new_run = OxmlElement('w:r')

    # Configurar propriedades do texto (cor azul, sublinhado)
    rPr = OxmlElement('w:rPr')

    # Cor azul
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    # Sublinhado
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    new_run.append(rPr)

    # Adicionar o texto
    new_run.text = texto_display
    hyperlink.append(new_run)

    return hyperlink

def processar_urls_em_paragrafo(paragraph):
    """
    VERSAO CORRIGIDA - Processa um paragrafo, convertendo URLs em hyperlinks
    mantendo a pontuacao original mas criando hyperlinks limpos
    """
    texto_completo = paragraph.text.strip()
    
    if not texto_completo:
        return False

    # Regex para encontrar URLs completas
    padrao_url = r'https?://[^\s<>"{}|\\^`\[\]]+(?:[^\s<>"{}|\\^`\[\]]*)'
    
    urls_encontradas = re.findall(padrao_url, texto_completo)

    if not urls_encontradas:
        return False

    # NOVA ABORDAGEM: Processar URLs mantendo a pontuacao original
    urls_processadas = []
    for url in urls_encontradas:
        # Separar a URL limpa da pontuacao
        url_limpa = re.sub(r'[),;.:!?]+$', '', url)
        pontuacao = url[len(url_limpa):] if len(url) > len(url_limpa) else ''
        
        if url_limpa and url_limpa not in [item[0] for item in urls_processadas]:
            urls_processadas.append((url_limpa, pontuacao, url))

    if not urls_processadas:
        return False

    print(f"   Encontradas {len(urls_processadas)} URLs: {[item[0] for item in urls_processadas[:2]]}{'...' if len(urls_processadas) > 2 else ''}")

    # Limpar o paragrafo atual
    _limpar_paragrafo(paragraph)

    # PROCESSAMENTO APRIMORADO: Manter pontuacao original
    texto_restante = texto_completo

    for url_limpa, pontuacao, url_original in urls_processadas:
        if url_original in texto_restante:
            # Dividir o texto pela URL original
            partes = texto_restante.split(url_original, 1)
            
            if len(partes) == 2:
                # Adicionar texto antes da URL (se houver)
                if partes[0]:
                    _add_text_preservando_asteriscos(paragraph, partes[0])

                # Criar hyperlink apenas com URL limpa
                hyperlink_element = adicionar_hyperlink(paragraph, url_limpa, url_limpa)
                paragraph._p.append(hyperlink_element)
                
                # Adicionar a pontuacao como texto normal (nao hyperlink)
                if pontuacao:
                    _add_text_preservando_asteriscos(paragraph, pontuacao)

                # Continuar com o resto do texto
                texto_restante = partes[1]

    # Adicionar texto restante apos a ultima URL (se houver)
    if texto_restante:
        _add_text_preservando_asteriscos(paragraph, texto_restante)

    return True

def converter_urls_docx_para_hyperlinks(arquivo_entrada, pasta_destino='/app/output', pasta_id_drive=None):
    # 1 Validar se o arquivo existe
    if not os.path.exists(arquivo_entrada):
        print(f"ERRO: Arquivo '{arquivo_entrada}' nao encontrado!")
        return False

    print(f"Abrindo arquivo: {arquivo_entrada}")
    doc = Document(arquivo_entrada)

    total_paragrafos_processados = 0
    total_urls_convertidas = 0

    # 2 Processar paragrafos principais
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            # Contar URLs no texto original antes de qualquer modificação
            texto_original = p.text
            urls_antes = len(re.findall(r'https?://[^\s]+', texto_original))

            # Primeiro tentar processar URLs (que pode reconstruir o parágrafo)
            if processar_urls_em_paragrafo(p):
                total_paragrafos_processados += 1
                total_urls_convertidas += urls_antes
                print(f"   Paragrafo {i+1} processado com {urls_antes} URLs")
            else:
                # Se não havia URLs, aplicar conversão de asteriscos normalmente
                converter_asteriscos_em_negrito(p)

    # 3 Processar tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        # Contar URLs no texto original antes de qualquer modificação
                        texto_original = p.text
                        urls_antes = len(re.findall(r'https?://[^\s]+', texto_original))

                        # Primeiro tentar processar URLs (pode reconstruir o parágrafo)
                        if processar_urls_em_paragrafo(p):
                            total_paragrafos_processados += 1
                            total_urls_convertidas += urls_antes
                        else:
                            # Se não havia URLs, aplicar conversão de asteriscos normalmente
                            converter_asteriscos_em_negrito(p)

    # 4 Gerar nome do arquivo final
    nome_base = os.path.basename(arquivo_entrada).replace('.docx', '')
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    arquivo_saida = os.path.join('output', f"{nome_base}_{timestamp}.docx")

    # 5 Salvar primeiro o arquivo localmente
    os.makedirs('output', exist_ok=True)
    doc.save(arquivo_saida)
    
    print(f"\nEstatisticas do processamento:")
    print(f"   - Paragrafos processados: {total_paragrafos_processados}")
    print(f"   - URLs convertidas em hyperlinks: {total_urls_convertidas}")
    print(f"Arquivo salvo localmente: {arquivo_saida}")

    # 6 Upload para Google Drive, se configurado
    if pasta_id_drive:
        try:
            upload_para_google_drive(arquivo_saida, os.path.basename(arquivo_saida), pasta_id_drive)
        except Exception as e:
            print(f"Erro no upload para Google Drive: {str(e)}")

    # 7 Copiar para a pasta compartilhada do Docker (opcional)
    if os.path.isdir(pasta_destino):
        destino_drive = os.path.join(pasta_destino, os.path.basename(arquivo_saida))
        try:
            # Evita SameFileError quando origem e destino sao o mesmo arquivo
            if not (os.path.exists(destino_drive) and os.path.samefile(arquivo_saida, destino_drive)):
                shutil.copy2(arquivo_saida, destino_drive)
                print(f"Arquivo tambem copiado para pasta compartilhada: {destino_drive}")
            else:
                print("Origem e destino sao o mesmo arquivo; copia ignorada.")
        except FileNotFoundError:
            # Alguns FS pedem que o diretorio exista antes do samefile; garanta e copie
            os.makedirs(pasta_destino, exist_ok=True)
            shutil.copy2(arquivo_saida, destino_drive)
            print(f"Arquivo tambem copiado para pasta compartilhada: {destino_drive}")
    else:
        print(f"Pasta destino '{pasta_destino}' nao encontrada. Pulei a copia local.")

    return True

def metodo_alternativo_melhorado(arquivo_entrada, arquivo_saida):
    """
    Metodo alternativo melhorado - substitui URLs por texto com formatacao
    """
    try:
        print(f"Metodo alternativo melhorado - Abrindo arquivo: {arquivo_entrada}")
        doc = Document(arquivo_entrada)

        total_urls_encontradas = 0
        paragrafos_processados = 0

        # Regex para capturar URLs completas
        padrao_url = r'https?://[^\s<>"{}|\\^`\[\]]+(?:[^\s<>"{}|\\^`\[\]]*)'

        # Processar paragrafos
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue

            texto_original = paragraph.text
            urls_no_texto = re.findall(padrao_url, texto_original)
            
            # Processar URLs mantendo pontuacao
            urls_processadas = []
            for url in urls_no_texto:
                url_limpa = re.sub(r'[),;.:!?]+$', '', url)
                pontuacao = url[len(url_limpa):] if len(url) > len(url_limpa) else ''
                
                if url_limpa and url_limpa not in [item[0] for item in urls_processadas]:
                    urls_processadas.append((url_limpa, pontuacao, url))

            if urls_processadas:
                print(f"   Paragrafo {paragrafos_processados + 1}: {len(urls_processadas)} URLs encontradas")
                total_urls_encontradas += len(urls_processadas)

                # Limpar o paragrafo
                _limpar_paragrafo(paragraph)

                # Reconstruir o paragrafo com formatacao
                texto_restante = texto_original

                for url_limpa, pontuacao, url_original in urls_processadas:
                    if url_original in texto_restante:
                        partes = texto_restante.split(url_original, 1)
                        
                        if len(partes) == 2:
                            # Adicionar texto antes da URL
                            if partes[0]:
                                paragraph.add_run(partes[0])

                            # Usar URL limpa no hyperlink
                            paragraph._p.append(adicionar_hyperlink(paragraph, url_limpa, url_limpa))
                            
                            # Adicionar pontuacao como texto normal
                            if pontuacao:
                                paragraph.add_run(pontuacao)

                            # Continuar com o resto
                            texto_restante = partes[1]

                # Adicionar texto restante
                if texto_restante:
                    paragraph.add_run(texto_restante)

            paragrafos_processados += 1

        # Processar tabelas tambem
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if not paragraph.text.strip():
                            continue

                        texto_original = paragraph.text
                        urls_no_texto = re.findall(padrao_url, texto_original)
                        
                        # Processar URLs mantendo pontuacao
                        urls_processadas = []
                        for url in urls_no_texto:
                            url_limpa = re.sub(r'[),;.:!?]+$', '', url)
                            pontuacao = url[len(url_limpa):] if len(url) > len(url_limpa) else ''
                            
                            if url_limpa and url_limpa not in [item[0] for item in urls_processadas]:
                                urls_processadas.append((url_limpa, pontuacao, url))

                        if urls_processadas:
                            total_urls_encontradas += len(urls_processadas)
                            _limpar_paragrafo(paragraph)

                            texto_restante = texto_original
                            for url_limpa, pontuacao, url_original in urls_processadas:
                                if url_original in texto_restante:
                                    partes = texto_restante.split(url_original, 1)
                                    
                                    if len(partes) == 2:
                                        if partes[0]:
                                            paragraph.add_run(partes[0])

                                        # Usar URL limpa no hyperlink
                                        paragraph._p.append(adicionar_hyperlink(paragraph, url_limpa, url_limpa))
                                        
                                        # Adicionar pontuacao como texto normal
                                        if pontuacao:
                                            paragraph.add_run(pontuacao)

                                        texto_restante = partes[1]

                            if texto_restante:
                                paragraph.add_run(texto_restante)

        # Salvar documento
        doc.save(arquivo_saida)

        # Upload para Google Drive
        arquivo_local = arquivo_saida
        nome_arquivo = os.path.basename(arquivo_saida)
        try:
            upload_para_google_drive(arquivo_local, nome_arquivo, "1HCo8W9Q9ak8aKOmMRPhSyVBntCS_GD6J")        
        except Exception as e:
            print(f"Erro no upload para Google Drive: {str(e)}")

        # Copiar para o Google Drive
        pasta_drive = r'/app/relatorios/'  # Altere para o caminho da sua pasta do Drive
        if os.path.isdir(pasta_drive):
            destino_drive = os.path.join(pasta_drive, os.path.basename(arquivo_saida))
            shutil.copy2(arquivo_saida, destino_drive)
            print(f"Arquivo tambem salvo em: {destino_drive}")
        else:
            print(f"Pasta do Google Drive nao encontrada: {pasta_drive}")
            
        print(f"\nMetodo alternativo concluido!")
        print(f"Estatisticas:")
        print(f"   - Paragrafos processados: {paragrafos_processados}")
        print(f"   - Total de URLs formatadas: {total_urls_encontradas}")
        print(f"Arquivo salvo como: {arquivo_saida}")

        return True

    except Exception as e:
        print(f"Erro no metodo alternativo: {str(e)}")
        return False

def testar_regex():
    """
    Funcao para testar a regex com URLs de exemplo
    """
    print("Testando regex com URLs de exemplo...")

    # URLs de teste
    urls_teste = [
        "https://tinyurl.com/2aymnjlf",
        "https://www.google.com/search?q=python",
        "http://example.com/path/to/file.html",
        "https://github.com/user/repo#readme",
        "https://site.com/page?param1=value1&param2=value2"
    ]

    # Regex melhorada
    padrao_url = r'https?://[^\s<>"{}|\\^`\[\]]+(?:[^\s<>"{}|\\^`\[\]]*)'

    for url in urls_teste:
        match = re.search(padrao_url, url)
        if match:
            print(f"   OK {url} -> Capturado: {match.group()}")
        else:
            print(f"   ERRO {url} -> Nao capturado")

    print("\n" + "="*50)

def gerar_versao_ajustada(arquivo_preliminar, pasta_id_drive=None):
    """
    Aplica os ajustes finais ao relatorio:
    - Converte URLs em hyperlinks
    - Gera nome do arquivo com timestamp
    - Salva localmente e realiza upload para o Google Drive (se configurado)
    """

    if not os.path.exists(arquivo_preliminar):
        print(f"Arquivo nao encontrado: {arquivo_preliminar}")
        return

    print(f"Aplicando versao ajustada com hyperlinks e timestamp...")
    
    # Reaproveitar funcao que ja processa os hyperlinks e salva com timestamp
    sucesso = converter_urls_docx_para_hyperlinks(arquivo_preliminar, pasta_id_drive=pasta_id_drive)

    if sucesso:
        print("Versao final ajustada com sucesso.")
    else:
        print("Falha ao gerar a versao ajustada.")