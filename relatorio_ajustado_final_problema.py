import os
import re
import shutil
from annotated_types import doc
from docx import Document
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google.oauth2 import service_account
from datetime import datetime
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def _limpar_paragrafo(paragraph):
    # remove todos os runs atuais
    for r in paragraph.runs[::-1]:
        paragraph._p.remove(r._element)

# Função para upload no Google Drive
def upload_para_google_drive(caminho_arquivo, nome_arquivo, pasta_id):
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload

    SERVICE_ACCOUNT_FILE = 'service_account.json'
    SCOPES = ['https://www.googleapis.com/auth/drive']

    creds = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES
    )

    service = build('drive', 'v3', credentials=creds)

    file_metadata = {
        'name': nome_arquivo,
        'parents': [pasta_id]
    }

    media = MediaFileUpload(caminho_arquivo, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    uploaded_file = service.files().create(
        body=file_metadata,
        media_body=media,
        fields='id',
        supportsAllDrives=True  # ✅ ESSENCIAL para pastas compartilhadas
    ).execute()

    print(f"✅ Upload concluído com ID: {uploaded_file.get('id')}")

def adicionar_hyperlink(paragraph, url, texto_display):
    """
    Adiciona um hyperlink a um parágrafo no documento Word
    """
    # Criar o elemento hyperlink
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), paragraph.part.relate_to(url,
                                                       "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                                                       is_external=True))

    # Criar o run com o texto do link
    new_run = OxmlElement('w:r')

    # Configurar propriedades do texto (cor azul, sublinhado)
    rPr = OxmlElement('w:rPr')

    # Cor azul
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    # Sublinhado
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    new_run.append(rPr)

    # Adicionar o texto
    new_run.text = texto_display
    hyperlink.append(new_run)

    return hyperlink

def processar_urls_em_paragrafo(paragraph):
    """
    VERSÃO CORRIGIDA - Processa um parágrafo, convertendo URLs em hyperlinks
    mantendo a pontuação original mas criando hyperlinks limpos
    """
    texto_completo = paragraph.text.strip()
    
    if not texto_completo:
        return False

    # ✅ REGEX para encontrar URLs completas
    padrao_url = r'https?://[^\s<>"{}|\\^`\[\]]+(?:[^\s<>"{}|\\^`\[\]]*)'
    
    urls_encontradas = re.findall(padrao_url, texto_completo)

    if not urls_encontradas:
        return False

    # ✅ NOVA ABORDAGEM: Processar URLs mantendo a pontuação original
    urls_processadas = []
    for url in urls_encontradas:
        # Separar a URL limpa da pontuação
        url_limpa = re.sub(r'[),;.:!?]+

def converter_urls_docx_para_hyperlinks(arquivo_entrada, pasta_destino='/app/output', pasta_id_drive=None):
    # 1️⃣ Validar se o arquivo existe
    if not os.path.exists(arquivo_entrada):
        print(f"❌ Erro: Arquivo '{arquivo_entrada}' não encontrado!")
        return False

    print(f"📖 Abrindo arquivo: {arquivo_entrada}")
    doc = Document(arquivo_entrada)

    total_paragrafos_processados = 0
    total_urls_convertidas = 0

    # 2️⃣ Processar parágrafos principais
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            urls_antes = len(re.findall(r'https?://[^\s]+', p.text))
            if processar_urls_em_paragrafo(p):
                total_paragrafos_processados += 1
                total_urls_convertidas += urls_antes
                print(f"   ✅ Parágrafo {i+1} processado com {urls_antes} URLs")

    # 3️⃣ Processar tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        urls_antes = len(re.findall(r'https?://[^\s]+', p.text))
                        if processar_urls_em_paragrafo(p):
                            total_paragrafos_processados += 1
                            total_urls_convertidas += urls_antes

    # 4️⃣ Gerar nome do arquivo final
    nome_base = os.path.basename(arquivo_entrada).replace('.docx', '')
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    arquivo_saida = os.path.join('output', f"{nome_base}_{timestamp}.docx")

    # 5️⃣ Salvar primeiro o arquivo localmente
    os.makedirs('output', exist_ok=True)
    doc.save(arquivo_saida)
    
    print(f"\n📊 Estatísticas do processamento:")
    print(f"   - Parágrafos processados: {total_paragrafos_processados}")
    print(f"   - URLs convertidas em hyperlinks: {total_urls_convertidas}")
    print(f"💾 Arquivo salvo localmente: {arquivo_saida}")

    # 6️⃣ Upload para Google Drive, se configurado
    if pasta_id_drive:
        try:
            upload_para_google_drive(arquivo_saida, os.path.basename(arquivo_saida), pasta_id_drive)
        except Exception as e:
            print(f"⚠️ Erro no upload para Google Drive: {str(e)}")

    # 7️⃣ Copiar para a pasta compartilhada do Docker (opcional)
    if os.path.isdir(pasta_destino):
        destino_drive = os.path.join(pasta_destino, os.path.basename(arquivo_saida))
        try:
            # Evita SameFileError quando origem e destino são o mesmo arquivo
            if not (os.path.exists(destino_drive) and os.path.samefile(arquivo_saida, destino_drive)):
                shutil.copy2(arquivo_saida, destino_drive)
                print(f"📂 Arquivo também copiado para pasta compartilhada: {destino_drive}")
            else:
                print("ℹ️ Origem e destino são o mesmo arquivo; cópia ignorada.")
        except FileNotFoundError:
            # Alguns FS pedem que o diretório exista antes do samefile; garanta e copie
            os.makedirs(pasta_destino, exist_ok=True)
            shutil.copy2(arquivo_saida, destino_drive)
            print(f"📂 Arquivo também copiado para pasta compartilhada: {destino_drive}")
    else:
        print(f"⚠️ Pasta destino '{pasta_destino}' não encontrada. Pulei a cópia local.")

    return True

def metodo_alternativo_melhorado(arquivo_entrada, arquivo_saida):
    """
    Método alternativo melhorado - substitui URLs por texto com formatação
    """
    try:
        print(f"📖 Método alternativo melhorado - Abrindo arquivo: {arquivo_entrada}")
        doc = Document(arquivo_entrada)

        total_urls_encontradas = 0
        paragrafos_processados = 0

        # ✅ Regex para capturar URLs completas
        padrao_url = r'https?://[^\s<>"{}|\\^`\[\]]+(?:[^\s<>"{}|\\^`\[\]]*)'

        # Processar parágrafos
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue

            texto_original = paragraph.text
            urls_no_texto = re.findall(padrao_url, texto_original)
            
            # ✅ Processar URLs mantendo pontuação
            urls_processadas = []
            for url in urls_no_texto:
                url_limpa = re.sub(r'[),;.:!?]+

            paragrafos_processados += 1

        # Processar tabelas também
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if not paragraph.text.strip():
                            continue

                        texto_original = paragraph.text
                        urls_no_texto = re.findall(padrao_url, texto_original)
                        
                        # ✅ Processar URLs mantendo pontuação
                        urls_processadas = []
                        for url in urls_no_texto:
                            url_limpa = re.sub(r'[),;.:!?]+

        # Salvar documento
        doc.save(arquivo_saida)

        # Upload para Google Drive
        arquivo_local = arquivo_saida
        nome_arquivo = os.path.basename(arquivo_saida)
        try:
            upload_para_google_drive(arquivo_local, nome_arquivo, "1HCo8W9Q9ak8aKOmMRPhSyVBntCS_GD6J")        
        except Exception as e:
            print(f"⚠️ Erro no upload para Google Drive: {str(e)}")

        # Copiar para o Google Drive
        pasta_drive = r'/app/relatorios/'  # Altere para o caminho da sua pasta do Drive
        if os.path.isdir(pasta_drive):
            destino_drive = os.path.join(pasta_drive, os.path.basename(arquivo_saida))
            shutil.copy2(arquivo_saida, destino_drive)
            print(f"📁 Arquivo também salvo em: {destino_drive}")
        else:
            print(f"⚠️ Pasta do Google Drive não encontrada: {pasta_drive}")
            
        print(f"\n✅ Método alternativo concluído!")
        print(f"📊 Estatísticas:")
        print(f"   - Parágrafos processados: {paragrafos_processados}")
        print(f"   - Total de URLs formatadas: {total_urls_encontradas}")
        print(f"💾 Arquivo salvo como: {arquivo_saida}")

        return True

    except Exception as e:
        print(f"❌ Erro no método alternativo: {str(e)}")
        return False

def testar_regex():
    """
    Função para testar a regex com URLs de exemplo
    """
    print("🧪 Testando regex com URLs de exemplo...")

    # URLs de teste
    urls_teste = [
        "https://tinyurl.com/2aymnjlf",
        "https://www.google.com/search?q=python",
        "http://example.com/path/to/file.html",
        "https://github.com/user/repo#readme",
        "https://site.com/page?param1=value1&param2=value2"
    ]

    # Regex melhorada
    padrao_url = r'https?://[^\s<>"{}|\\^`\[\]\(\),;]+(?:[^\s<>"{}|\\^`\[\]\(\),;.])*'

    for url in urls_teste:
        match = re.search(padrao_url, url)
        if match:
            print(f"   ✅ {url} -> Capturado: {match.group()}")
        else:
            print(f"   ❌ {url} -> Não capturado")

    print("\n" + "="*50)

def gerar_versao_ajustada(arquivo_preliminar, pasta_id_drive=None):
    """
    Aplica os ajustes finais ao relatório:
    - Converte URLs em hyperlinks
    - Gera nome do arquivo com timestamp
    - Salva localmente e realiza upload para o Google Drive (se configurado)
    """

    if not os.path.exists(arquivo_preliminar):
        print(f"❌ Arquivo não encontrado: {arquivo_preliminar}")
        return

    print(f"📖 Aplicando versão ajustada com hyperlinks e timestamp...")
    
    # 🧠 Reaproveitar função que já processa os hyperlinks e salva com timestamp
    sucesso = converter_urls_docx_para_hyperlinks(arquivo_preliminar, pasta_id_drive=pasta_id_drive)

    if sucesso:
        print("✅ Versão final ajustada com sucesso.")
    else:
        print("❌ Falha ao gerar a versão ajustada."), '', url)
        pontuacao = url[len(url_limpa):] if len(url) > len(url_limpa) else ''
        
        if url_limpa and url_limpa not in [item[0] for item in urls_processadas]:
            urls_processadas.append((url_limpa, pontuacao, url))

    if not urls_processadas:
        return False

    print(f"   🔗 Encontradas {len(urls_processadas)} URLs: {[item[0] for item in urls_processadas[:2]]}{'...' if len(urls_processadas) > 2 else ''}")

    # Limpar o parágrafo atual
    _limpar_paragrafo(paragraph)

    # ✅ PROCESSAMENTO APRIMORADO: Manter pontuação original
    texto_restante = texto_completo

    for url_limpa, pontuacao, url_original in urls_processadas:
        if url_original in texto_restante:
            # Dividir o texto pela URL original
            partes = texto_restante.split(url_original, 1)
            
            if len(partes) == 2:
                # Adicionar texto antes da URL (se houver)
                if partes[0]:
                    paragraph.add_run(partes[0])

                # ✅ Criar hyperlink apenas com URL limpa
                hyperlink_element = adicionar_hyperlink(paragraph, url_limpa, url_limpa)
                paragraph._p.append(hyperlink_element)
                
                # ✅ Adicionar a pontuação como texto normal (não hyperlink)
                if pontuacao:
                    paragraph.add_run(pontuacao)

                # Continuar com o resto do texto
                texto_restante = partes[1]

    # Adicionar texto restante após a última URL (se houver)
    if texto_restante:
        paragraph.add_run(texto_restante)

    return True

def converter_urls_docx_para_hyperlinks(arquivo_entrada, pasta_destino='/app/output', pasta_id_drive=None):
    # 1️⃣ Validar se o arquivo existe
    if not os.path.exists(arquivo_entrada):
        print(f"❌ Erro: Arquivo '{arquivo_entrada}' não encontrado!")
        return False

    print(f"📖 Abrindo arquivo: {arquivo_entrada}")
    doc = Document(arquivo_entrada)

    total_paragrafos_processados = 0
    total_urls_convertidas = 0

    # 2️⃣ Processar parágrafos principais
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            urls_antes = len(re.findall(r'https?://[^\s]+', p.text))
            if processar_urls_em_paragrafo(p):
                total_paragrafos_processados += 1
                total_urls_convertidas += urls_antes
                print(f"   ✅ Parágrafo {i+1} processado com {urls_antes} URLs")

    # 3️⃣ Processar tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        urls_antes = len(re.findall(r'https?://[^\s]+', p.text))
                        if processar_urls_em_paragrafo(p):
                            total_paragrafos_processados += 1
                            total_urls_convertidas += urls_antes

    # 4️⃣ Gerar nome do arquivo final
    nome_base = os.path.basename(arquivo_entrada).replace('.docx', '')
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    arquivo_saida = os.path.join('output', f"{nome_base}_{timestamp}.docx")

    # 5️⃣ Salvar primeiro o arquivo localmente
    os.makedirs('output', exist_ok=True)
    doc.save(arquivo_saida)
    
    print(f"\n📊 Estatísticas do processamento:")
    print(f"   - Parágrafos processados: {total_paragrafos_processados}")
    print(f"   - URLs convertidas em hyperlinks: {total_urls_convertidas}")
    print(f"💾 Arquivo salvo localmente: {arquivo_saida}")

    # 6️⃣ Upload para Google Drive, se configurado
    if pasta_id_drive:
        try:
            upload_para_google_drive(arquivo_saida, os.path.basename(arquivo_saida), pasta_id_drive)
        except Exception as e:
            print(f"⚠️ Erro no upload para Google Drive: {str(e)}")

    # 7️⃣ Copiar para a pasta compartilhada do Docker (opcional)
    if os.path.isdir(pasta_destino):
        destino_drive = os.path.join(pasta_destino, os.path.basename(arquivo_saida))
        try:
            # Evita SameFileError quando origem e destino são o mesmo arquivo
            if not (os.path.exists(destino_drive) and os.path.samefile(arquivo_saida, destino_drive)):
                shutil.copy2(arquivo_saida, destino_drive)
                print(f"📂 Arquivo também copiado para pasta compartilhada: {destino_drive}")
            else:
                print("ℹ️ Origem e destino são o mesmo arquivo; cópia ignorada.")
        except FileNotFoundError:
            # Alguns FS pedem que o diretório exista antes do samefile; garanta e copie
            os.makedirs(pasta_destino, exist_ok=True)
            shutil.copy2(arquivo_saida, destino_drive)
            print(f"📂 Arquivo também copiado para pasta compartilhada: {destino_drive}")
    else:
        print(f"⚠️ Pasta destino '{pasta_destino}' não encontrada. Pulei a cópia local.")

    return True

def metodo_alternativo_melhorado(arquivo_entrada, arquivo_saida):
    """
    Método alternativo melhorado - substitui URLs por texto com formatação
    """
    try:
        print(f"📖 Método alternativo melhorado - Abrindo arquivo: {arquivo_entrada}")
        doc = Document(arquivo_entrada)

        total_urls_encontradas = 0
        paragrafos_processados = 0

        # ✅ Regex corrigida para capturar URLs completas sem caracteres indesejados
        padrao_url = r'https?://[^\s<>"{}|\\^`\[\]\(\),;]+(?:[^\s<>"{}|\\^`\[\]\(\),;.])*'

        # Processar parágrafos
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue

            texto_original = paragraph.text
            urls_no_texto = re.findall(padrao_url, texto_original)
            
            # ✅ Limpeza das URLs encontradas
            urls_limpas = []
            for url in urls_no_texto:
                url_limpa = re.sub(r'[),;.:!?]+$', '', url)
                if url_limpa and url_limpa not in urls_limpas:
                    urls_limpas.append(url_limpa)

            if urls_limpas:
                print(f"   🔗 Parágrafo {paragrafos_processados + 1}: {len(urls_limpas)} URLs encontradas")
                total_urls_encontradas += len(urls_limpas)

                # Limpar o parágrafo
                _limpar_paragrafo(paragraph)

                # Reconstruir o parágrafo com formatação
                texto_restante = texto_original

                for url_limpa in urls_limpas:
                    # Procurar pela URL original no texto
                    padrao_busca = re.escape(url_limpa) + r'[),;.:!?]*'
                    match = re.search(padrao_busca, texto_restante)
                    
                    if match:
                        url_original = match.group()
                        partes = texto_restante.split(url_original, 1)
                        
                        if len(partes) == 2:
                            # Adicionar texto antes da URL
                            if partes[0]:
                                paragraph.add_run(partes[0])

                            # ✅ Usar URL limpa
                            paragraph._p.append(adicionar_hyperlink(paragraph, url_limpa, url_limpa))

                            # Continuar com o resto
                            texto_restante = partes[1]

                # Adicionar texto restante
                if texto_restante:
                    paragraph.add_run(texto_restante)

            paragrafos_processados += 1

        # Processar tabelas também
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if not paragraph.text.strip():
                            continue

                        texto_original = paragraph.text
                        urls_no_texto = re.findall(padrao_url, texto_original)
                        
                        # ✅ Limpeza das URLs encontradas
                        urls_limpas = []
                        for url in urls_no_texto:
                            url_limpa = re.sub(r'[),;.:!?]+$', '', url)
                            if url_limpa and url_limpa not in urls_limpas:
                                urls_limpas.append(url_limpa)

                        if urls_limpas:
                            total_urls_encontradas += len(urls_limpas)
                            _limpar_paragrafo(paragraph)

                            texto_restante = texto_original
                            for url_limpa in urls_limpas:
                                padrao_busca = re.escape(url_limpa) + r'[),;.:!?]*'
                                match = re.search(padrao_busca, texto_restante)
                                
                                if match:
                                    url_original = match.group()
                                    partes = texto_restante.split(url_original, 1)
                                    
                                    if len(partes) == 2:
                                        if partes[0]:
                                            paragraph.add_run(partes[0])

                                        # ✅ Usar URL limpa
                                        paragraph._p.append(adicionar_hyperlink(paragraph, url_limpa, url_limpa))

                                        texto_restante = partes[1]

                            if texto_restante:
                                paragraph.add_run(texto_restante)

        # Salvar documento
        doc.save(arquivo_saida)

        # Upload para Google Drive
        arquivo_local = arquivo_saida
        nome_arquivo = os.path.basename(arquivo_saida)
        try:
            upload_para_google_drive(arquivo_local, nome_arquivo, "1HCo8W9Q9ak8aKOmMRPhSyVBntCS_GD6J")        
        except Exception as e:
            print(f"⚠️ Erro no upload para Google Drive: {str(e)}")

        # Copiar para o Google Drive
        pasta_drive = r'/app/relatorios/'  # Altere para o caminho da sua pasta do Drive
        if os.path.isdir(pasta_drive):
            destino_drive = os.path.join(pasta_drive, os.path.basename(arquivo_saida))
            shutil.copy2(arquivo_saida, destino_drive)
            print(f"📁 Arquivo também salvo em: {destino_drive}")
        else:
            print(f"⚠️ Pasta do Google Drive não encontrada: {pasta_drive}")
            
        print(f"\n✅ Método alternativo concluído!")
        print(f"📊 Estatísticas:")
        print(f"   - Parágrafos processados: {paragrafos_processados}")
        print(f"   - Total de URLs formatadas: {total_urls_encontradas}")
        print(f"💾 Arquivo salvo como: {arquivo_saida}")

        return True

    except Exception as e:
        print(f"❌ Erro no método alternativo: {str(e)}")
        return False

def testar_regex():
    """
    Função para testar a regex com URLs de exemplo
    """
    print("🧪 Testando regex com URLs de exemplo...")

    # URLs de teste
    urls_teste = [
        "https://tinyurl.com/2aymnjlf",
        "https://www.google.com/search?q=python",
        "http://example.com/path/to/file.html",
        "https://github.com/user/repo#readme",
        "https://site.com/page?param1=value1&param2=value2"
    ]

    # Regex melhorada
    padrao_url = r'https?://[^\s<>"{}|\\^`\[\]\(\),;]+(?:[^\s<>"{}|\\^`\[\]\(\),;.])*'

    for url in urls_teste:
        match = re.search(padrao_url, url)
        if match:
            print(f"   ✅ {url} -> Capturado: {match.group()}")
        else:
            print(f"   ❌ {url} -> Não capturado")

    print("\n" + "="*50)

def gerar_versao_ajustada(arquivo_preliminar, pasta_id_drive=None):
    """
    Aplica os ajustes finais ao relatório:
    - Converte URLs em hyperlinks
    - Gera nome do arquivo com timestamp
    - Salva localmente e realiza upload para o Google Drive (se configurado)
    """

    if not os.path.exists(arquivo_preliminar):
        print(f"❌ Arquivo não encontrado: {arquivo_preliminar}")
        return

    print(f"📖 Aplicando versão ajustada com hyperlinks e timestamp...")
    
    # 🧠 Reaproveitar função que já processa os hyperlinks e salva com timestamp
    sucesso = converter_urls_docx_para_hyperlinks(arquivo_preliminar, pasta_id_drive=pasta_id_drive)

    if sucesso:
        print("✅ Versão final ajustada com sucesso.")
    else:
        print("❌ Falha ao gerar a versão ajustada."), '', url)
                pontuacao = url[len(url_limpa):] if len(url) > len(url_limpa) else ''
                
                if url_limpa and url_limpa not in [item[0] for item in urls_processadas]:
                    urls_processadas.append((url_limpa, pontuacao, url))

            if urls_processadas:
                print(f"   🔗 Parágrafo {paragrafos_processados + 1}: {len(urls_processadas)} URLs encontradas")
                total_urls_encontradas += len(urls_processadas)

                # Limpar o parágrafo
                _limpar_paragrafo(paragraph)

                # Reconstruir o parágrafo com formatação
                texto_restante = texto_original

                for url_limpa, pontuacao, url_original in urls_processadas:
                    if url_original in texto_restante:
                        partes = texto_restante.split(url_original, 1)
                        
                        if len(partes) == 2:
                            # Adicionar texto antes da URL
                            if partes[0]:
                                paragraph.add_run(partes[0])

                            # ✅ Usar URL limpa no hyperlink
                            paragraph._p.append(adicionar_hyperlink(paragraph, url_limpa, url_limpa))
                            
                            # ✅ Adicionar pontuação como texto normal
                            if pontuacao:
                                paragraph.add_run(pontuacao)

                            # Continuar com o resto
                            texto_restante = partes[1]

                # Adicionar texto restante
                if texto_restante:
                    paragraph.add_run(texto_restante)

            paragrafos_processados += 1

        # Processar tabelas também
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if not paragraph.text.strip():
                            continue

                        texto_original = paragraph.text
                        urls_no_texto = re.findall(padrao_url, texto_original)
                        
                        # ✅ Limpeza das URLs encontradas
                        urls_limpas = []
                        for url in urls_no_texto:
                            url_limpa = re.sub(r'[),;.:!?]+$', '', url)
                            if url_limpa and url_limpa not in urls_limpas:
                                urls_limpas.append(url_limpa)

                        if urls_limpas:
                            total_urls_encontradas += len(urls_limpas)
                            _limpar_paragrafo(paragraph)

                            texto_restante = texto_original
                            for url_limpa in urls_limpas:
                                padrao_busca = re.escape(url_limpa) + r'[),;.:!?]*'
                                match = re.search(padrao_busca, texto_restante)
                                
                                if match:
                                    url_original = match.group()
                                    partes = texto_restante.split(url_original, 1)
                                    
                                    if len(partes) == 2:
                                        if partes[0]:
                                            paragraph.add_run(partes[0])

                                        # ✅ Usar URL limpa
                                        paragraph._p.append(adicionar_hyperlink(paragraph, url_limpa, url_limpa))

                                        texto_restante = partes[1]

                            if texto_restante:
                                paragraph.add_run(texto_restante)

        # Salvar documento
        doc.save(arquivo_saida)

        # Upload para Google Drive
        arquivo_local = arquivo_saida
        nome_arquivo = os.path.basename(arquivo_saida)
        try:
            upload_para_google_drive(arquivo_local, nome_arquivo, "1HCo8W9Q9ak8aKOmMRPhSyVBntCS_GD6J")        
        except Exception as e:
            print(f"⚠️ Erro no upload para Google Drive: {str(e)}")

        # Copiar para o Google Drive
        pasta_drive = r'/app/relatorios/'  # Altere para o caminho da sua pasta do Drive
        if os.path.isdir(pasta_drive):
            destino_drive = os.path.join(pasta_drive, os.path.basename(arquivo_saida))
            shutil.copy2(arquivo_saida, destino_drive)
            print(f"📁 Arquivo também salvo em: {destino_drive}")
        else:
            print(f"⚠️ Pasta do Google Drive não encontrada: {pasta_drive}")
            
        print(f"\n✅ Método alternativo concluído!")
        print(f"📊 Estatísticas:")
        print(f"   - Parágrafos processados: {paragrafos_processados}")
        print(f"   - Total de URLs formatadas: {total_urls_encontradas}")
        print(f"💾 Arquivo salvo como: {arquivo_saida}")

        return True

    except Exception as e:
        print(f"❌ Erro no método alternativo: {str(e)}")
        return False

def testar_regex():
    """
    Função para testar a regex com URLs de exemplo
    """
    print("🧪 Testando regex com URLs de exemplo...")

    # URLs de teste
    urls_teste = [
        "https://tinyurl.com/2aymnjlf",
        "https://www.google.com/search?q=python",
        "http://example.com/path/to/file.html",
        "https://github.com/user/repo#readme",
        "https://site.com/page?param1=value1&param2=value2"
    ]

    # Regex melhorada
    padrao_url = r'https?://[^\s<>"{}|\\^`\[\]\(\),;]+(?:[^\s<>"{}|\\^`\[\]\(\),;.])*'

    for url in urls_teste:
        match = re.search(padrao_url, url)
        if match:
            print(f"   ✅ {url} -> Capturado: {match.group()}")
        else:
            print(f"   ❌ {url} -> Não capturado")

    print("\n" + "="*50)

def gerar_versao_ajustada(arquivo_preliminar, pasta_id_drive=None):
    """
    Aplica os ajustes finais ao relatório:
    - Converte URLs em hyperlinks
    - Gera nome do arquivo com timestamp
    - Salva localmente e realiza upload para o Google Drive (se configurado)
    """

    if not os.path.exists(arquivo_preliminar):
        print(f"❌ Arquivo não encontrado: {arquivo_preliminar}")
        return

    print(f"📖 Aplicando versão ajustada com hyperlinks e timestamp...")
    
    # 🧠 Reaproveitar função que já processa os hyperlinks e salva com timestamp
    sucesso = converter_urls_docx_para_hyperlinks(arquivo_preliminar, pasta_id_drive=pasta_id_drive)

    if sucesso:
        print("✅ Versão final ajustada com sucesso.")
    else:
        print("❌ Falha ao gerar a versão ajustada."), '', url)
        pontuacao = url[len(url_limpa):] if len(url) > len(url_limpa) else ''
        
        if url_limpa and url_limpa not in [item[0] for item in urls_processadas]:
            urls_processadas.append((url_limpa, pontuacao, url))

    if not urls_processadas:
        return False

    print(f"   🔗 Encontradas {len(urls_processadas)} URLs: {[item[0] for item in urls_processadas[:2]]}{'...' if len(urls_processadas) > 2 else ''}")

    # Limpar o parágrafo atual
    _limpar_paragrafo(paragraph)

    # ✅ PROCESSAMENTO APRIMORADO: Manter pontuação original
    texto_restante = texto_completo

    for url_limpa, pontuacao, url_original in urls_processadas:
        if url_original in texto_restante:
            # Dividir o texto pela URL original
            partes = texto_restante.split(url_original, 1)
            
            if len(partes) == 2:
                # Adicionar texto antes da URL (se houver)
                if partes[0]:
                    paragraph.add_run(partes[0])

                # ✅ Criar hyperlink apenas com URL limpa
                hyperlink_element = adicionar_hyperlink(paragraph, url_limpa, url_limpa)
                paragraph._p.append(hyperlink_element)
                
                # ✅ Adicionar a pontuação como texto normal (não hyperlink)
                if pontuacao:
                    paragraph.add_run(pontuacao)

                # Continuar com o resto do texto
                texto_restante = partes[1]

    # Adicionar texto restante após a última URL (se houver)
    if texto_restante:
        paragraph.add_run(texto_restante)

    return True

def converter_urls_docx_para_hyperlinks(arquivo_entrada, pasta_destino='/app/output', pasta_id_drive=None):
    # 1️⃣ Validar se o arquivo existe
    if not os.path.exists(arquivo_entrada):
        print(f"❌ Erro: Arquivo '{arquivo_entrada}' não encontrado!")
        return False

    print(f"📖 Abrindo arquivo: {arquivo_entrada}")
    doc = Document(arquivo_entrada)

    total_paragrafos_processados = 0
    total_urls_convertidas = 0

    # 2️⃣ Processar parágrafos principais
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            urls_antes = len(re.findall(r'https?://[^\s]+', p.text))
            if processar_urls_em_paragrafo(p):
                total_paragrafos_processados += 1
                total_urls_convertidas += urls_antes
                print(f"   ✅ Parágrafo {i+1} processado com {urls_antes} URLs")

    # 3️⃣ Processar tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        urls_antes = len(re.findall(r'https?://[^\s]+', p.text))
                        if processar_urls_em_paragrafo(p):
                            total_paragrafos_processados += 1
                            total_urls_convertidas += urls_antes

    # 4️⃣ Gerar nome do arquivo final
    nome_base = os.path.basename(arquivo_entrada).replace('.docx', '')
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    arquivo_saida = os.path.join('output', f"{nome_base}_{timestamp}.docx")

    # 5️⃣ Salvar primeiro o arquivo localmente
    os.makedirs('output', exist_ok=True)
    doc.save(arquivo_saida)
    
    print(f"\n📊 Estatísticas do processamento:")
    print(f"   - Parágrafos processados: {total_paragrafos_processados}")
    print(f"   - URLs convertidas em hyperlinks: {total_urls_convertidas}")
    print(f"💾 Arquivo salvo localmente: {arquivo_saida}")

    # 6️⃣ Upload para Google Drive, se configurado
    if pasta_id_drive:
        try:
            upload_para_google_drive(arquivo_saida, os.path.basename(arquivo_saida), pasta_id_drive)
        except Exception as e:
            print(f"⚠️ Erro no upload para Google Drive: {str(e)}")

    # 7️⃣ Copiar para a pasta compartilhada do Docker (opcional)
    if os.path.isdir(pasta_destino):
        destino_drive = os.path.join(pasta_destino, os.path.basename(arquivo_saida))
        try:
            # Evita SameFileError quando origem e destino são o mesmo arquivo
            if not (os.path.exists(destino_drive) and os.path.samefile(arquivo_saida, destino_drive)):
                shutil.copy2(arquivo_saida, destino_drive)
                print(f"📂 Arquivo também copiado para pasta compartilhada: {destino_drive}")
            else:
                print("ℹ️ Origem e destino são o mesmo arquivo; cópia ignorada.")
        except FileNotFoundError:
            # Alguns FS pedem que o diretório exista antes do samefile; garanta e copie
            os.makedirs(pasta_destino, exist_ok=True)
            shutil.copy2(arquivo_saida, destino_drive)
            print(f"📂 Arquivo também copiado para pasta compartilhada: {destino_drive}")
    else:
        print(f"⚠️ Pasta destino '{pasta_destino}' não encontrada. Pulei a cópia local.")

    return True

def metodo_alternativo_melhorado(arquivo_entrada, arquivo_saida):
    """
    Método alternativo melhorado - substitui URLs por texto com formatação
    """
    try:
        print(f"📖 Método alternativo melhorado - Abrindo arquivo: {arquivo_entrada}")
        doc = Document(arquivo_entrada)

        total_urls_encontradas = 0
        paragrafos_processados = 0

        # ✅ Regex corrigida para capturar URLs completas sem caracteres indesejados
        padrao_url = r'https?://[^\s<>"{}|\\^`\[\]\(\),;]+(?:[^\s<>"{}|\\^`\[\]\(\),;.])*'

        # Processar parágrafos
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue

            texto_original = paragraph.text
            urls_no_texto = re.findall(padrao_url, texto_original)
            
            # ✅ Limpeza das URLs encontradas
            urls_limpas = []
            for url in urls_no_texto:
                url_limpa = re.sub(r'[),;.:!?]+$', '', url)
                if url_limpa and url_limpa not in urls_limpas:
                    urls_limpas.append(url_limpa)

            if urls_limpas:
                print(f"   🔗 Parágrafo {paragrafos_processados + 1}: {len(urls_limpas)} URLs encontradas")
                total_urls_encontradas += len(urls_limpas)

                # Limpar o parágrafo
                _limpar_paragrafo(paragraph)

                # Reconstruir o parágrafo com formatação
                texto_restante = texto_original

                for url_limpa in urls_limpas:
                    # Procurar pela URL original no texto
                    padrao_busca = re.escape(url_limpa) + r'[),;.:!?]*'
                    match = re.search(padrao_busca, texto_restante)
                    
                    if match:
                        url_original = match.group()
                        partes = texto_restante.split(url_original, 1)
                        
                        if len(partes) == 2:
                            # Adicionar texto antes da URL
                            if partes[0]:
                                paragraph.add_run(partes[0])

                            # ✅ Usar URL limpa
                            paragraph._p.append(adicionar_hyperlink(paragraph, url_limpa, url_limpa))

                            # Continuar com o resto
                            texto_restante = partes[1]

                # Adicionar texto restante
                if texto_restante:
                    paragraph.add_run(texto_restante)

            paragrafos_processados += 1

        # Processar tabelas também
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if not paragraph.text.strip():
                            continue

                        texto_original = paragraph.text
                        urls_no_texto = re.findall(padrao_url, texto_original)
                        
                        # ✅ Limpeza das URLs encontradas
                        urls_limpas = []
                        for url in urls_no_texto:
                            url_limpa = re.sub(r'[),;.:!?]+$', '', url)
                            if url_limpa and url_limpa not in urls_limpas:
                                urls_limpas.append(url_limpa)

                        if urls_limpas:
                            total_urls_encontradas += len(urls_limpas)
                            _limpar_paragrafo(paragraph)

                            texto_restante = texto_original
                            for url_limpa in urls_limpas:
                                padrao_busca = re.escape(url_limpa) + r'[),;.:!?]*'
                                match = re.search(padrao_busca, texto_restante)
                                
                                if match:
                                    url_original = match.group()
                                    partes = texto_restante.split(url_original, 1)
                                    
                                    if len(partes) == 2:
                                        if partes[0]:
                                            paragraph.add_run(partes[0])

                                        # ✅ Usar URL limpa
                                        paragraph._p.append(adicionar_hyperlink(paragraph, url_limpa, url_limpa))

                                        texto_restante = partes[1]

                            if texto_restante:
                                paragraph.add_run(texto_restante)

        # Salvar documento
        doc.save(arquivo_saida)

        # Upload para Google Drive
        arquivo_local = arquivo_saida
        nome_arquivo = os.path.basename(arquivo_saida)
        try:
            upload_para_google_drive(arquivo_local, nome_arquivo, "1HCo8W9Q9ak8aKOmMRPhSyVBntCS_GD6J")        
        except Exception as e:
            print(f"⚠️ Erro no upload para Google Drive: {str(e)}")

        # Copiar para o Google Drive
        pasta_drive = r'/app/relatorios/'  # Altere para o caminho da sua pasta do Drive
        if os.path.isdir(pasta_drive):
            destino_drive = os.path.join(pasta_drive, os.path.basename(arquivo_saida))
            shutil.copy2(arquivo_saida, destino_drive)
            print(f"📁 Arquivo também salvo em: {destino_drive}")
        else:
            print(f"⚠️ Pasta do Google Drive não encontrada: {pasta_drive}")
            
        print(f"\n✅ Método alternativo concluído!")
        print(f"📊 Estatísticas:")
        print(f"   - Parágrafos processados: {paragrafos_processados}")
        print(f"   - Total de URLs formatadas: {total_urls_encontradas}")
        print(f"💾 Arquivo salvo como: {arquivo_saida}")

        return True

    except Exception as e:
        print(f"❌ Erro no método alternativo: {str(e)}")
        return False

def testar_regex():
    """
    Função para testar a regex com URLs de exemplo
    """
    print("🧪 Testando regex com URLs de exemplo...")

    # URLs de teste
    urls_teste = [
        "https://tinyurl.com/2aymnjlf",
        "https://www.google.com/search?q=python",
        "http://example.com/path/to/file.html",
        "https://github.com/user/repo#readme",
        "https://site.com/page?param1=value1&param2=value2"
    ]

    # Regex melhorada
    padrao_url = r'https?://[^\s<>"{}|\\^`\[\]\(\),;]+(?:[^\s<>"{}|\\^`\[\]\(\),;.])*'

    for url in urls_teste:
        match = re.search(padrao_url, url)
        if match:
            print(f"   ✅ {url} -> Capturado: {match.group()}")
        else:
            print(f"   ❌ {url} -> Não capturado")

    print("\n" + "="*50)

def gerar_versao_ajustada(arquivo_preliminar, pasta_id_drive=None):
    """
    Aplica os ajustes finais ao relatório:
    - Converte URLs em hyperlinks
    - Gera nome do arquivo com timestamp
    - Salva localmente e realiza upload para o Google Drive (se configurado)
    """

    if not os.path.exists(arquivo_preliminar):
        print(f"❌ Arquivo não encontrado: {arquivo_preliminar}")
        return

    print(f"📖 Aplicando versão ajustada com hyperlinks e timestamp...")
    
    # 🧠 Reaproveitar função que já processa os hyperlinks e salva com timestamp
    sucesso = converter_urls_docx_para_hyperlinks(arquivo_preliminar, pasta_id_drive=pasta_id_drive)

    if sucesso:
        print("✅ Versão final ajustada com sucesso.")
    else:
        print("❌ Falha ao gerar a versão ajustada."), '', url)
                            pontuacao = url[len(url_limpa):] if len(url) > len(url_limpa) else ''
                            
                            if url_limpa and url_limpa not in [item[0] for item in urls_processadas]:
                                urls_processadas.append((url_limpa, pontuacao, url))

                        if urls_processadas:
                            total_urls_encontradas += len(urls_processadas)
                            _limpar_paragrafo(paragraph)

                            texto_restante = texto_original
                            for url_limpa, pontuacao, url_original in urls_processadas:
                                if url_original in texto_restante:
                                    partes = texto_restante.split(url_original, 1)
                                    
                                    if len(partes) == 2:
                                        if partes[0]:
                                            paragraph.add_run(partes[0])

                                        # ✅ Usar URL limpa no hyperlink
                                        paragraph._p.append(adicionar_hyperlink(paragraph, url_limpa, url_limpa))
                                        
                                        # ✅ Adicionar pontuação como texto normal
                                        if pontuacao:
                                            paragraph.add_run(pontuacao)

                                        texto_restante = partes[1]

                            if texto_restante:
                                paragraph.add_run(texto_restante)

        # Salvar documento
        doc.save(arquivo_saida)

        # Upload para Google Drive
        arquivo_local = arquivo_saida
        nome_arquivo = os.path.basename(arquivo_saida)
        try:
            upload_para_google_drive(arquivo_local, nome_arquivo, "1HCo8W9Q9ak8aKOmMRPhSyVBntCS_GD6J")        
        except Exception as e:
            print(f"⚠️ Erro no upload para Google Drive: {str(e)}")

        # Copiar para o Google Drive
        pasta_drive = r'/app/relatorios/'  # Altere para o caminho da sua pasta do Drive
        if os.path.isdir(pasta_drive):
            destino_drive = os.path.join(pasta_drive, os.path.basename(arquivo_saida))
            shutil.copy2(arquivo_saida, destino_drive)
            print(f"📁 Arquivo também salvo em: {destino_drive}")
        else:
            print(f"⚠️ Pasta do Google Drive não encontrada: {pasta_drive}")
            
        print(f"\n✅ Método alternativo concluído!")
        print(f"📊 Estatísticas:")
        print(f"   - Parágrafos processados: {paragrafos_processados}")
        print(f"   - Total de URLs formatadas: {total_urls_encontradas}")
        print(f"💾 Arquivo salvo como: {arquivo_saida}")

        return True

    except Exception as e:
        print(f"❌ Erro no método alternativo: {str(e)}")
        return False

def testar_regex():
    """
    Função para testar a regex com URLs de exemplo
    """
    print("🧪 Testando regex com URLs de exemplo...")

    # URLs de teste
    urls_teste = [
        "https://tinyurl.com/2aymnjlf",
        "https://www.google.com/search?q=python",
        "http://example.com/path/to/file.html",
        "https://github.com/user/repo#readme",
        "https://site.com/page?param1=value1&param2=value2"
    ]

    # Regex melhorada
    padrao_url = r'https?://[^\s<>"{}|\\^`\[\]\(\),;]+(?:[^\s<>"{}|\\^`\[\]\(\),;.])*'

    for url in urls_teste:
        match = re.search(padrao_url, url)
        if match:
            print(f"   ✅ {url} -> Capturado: {match.group()}")
        else:
            print(f"   ❌ {url} -> Não capturado")

    print("\n" + "="*50)

def gerar_versao_ajustada(arquivo_preliminar, pasta_id_drive=None):
    """
    Aplica os ajustes finais ao relatório:
    - Converte URLs em hyperlinks
    - Gera nome do arquivo com timestamp
    - Salva localmente e realiza upload para o Google Drive (se configurado)
    """

    if not os.path.exists(arquivo_preliminar):
        print(f"❌ Arquivo não encontrado: {arquivo_preliminar}")
        return

    print(f"📖 Aplicando versão ajustada com hyperlinks e timestamp...")
    
    # 🧠 Reaproveitar função que já processa os hyperlinks e salva com timestamp
    sucesso = converter_urls_docx_para_hyperlinks(arquivo_preliminar, pasta_id_drive=pasta_id_drive)

    if sucesso:
        print("✅ Versão final ajustada com sucesso.")
    else:
        print("❌ Falha ao gerar a versão ajustada."), '', url)
        pontuacao = url[len(url_limpa):] if len(url) > len(url_limpa) else ''
        
        if url_limpa and url_limpa not in [item[0] for item in urls_processadas]:
            urls_processadas.append((url_limpa, pontuacao, url))

    if not urls_processadas:
        return False

    print(f"   🔗 Encontradas {len(urls_processadas)} URLs: {[item[0] for item in urls_processadas[:2]]}{'...' if len(urls_processadas) > 2 else ''}")

    # Limpar o parágrafo atual
    _limpar_paragrafo(paragraph)

    # ✅ PROCESSAMENTO APRIMORADO: Manter pontuação original
    texto_restante = texto_completo

    for url_limpa, pontuacao, url_original in urls_processadas:
        if url_original in texto_restante:
            # Dividir o texto pela URL original
            partes = texto_restante.split(url_original, 1)
            
            if len(partes) == 2:
                # Adicionar texto antes da URL (se houver)
                if partes[0]:
                    paragraph.add_run(partes[0])

                # ✅ Criar hyperlink apenas com URL limpa
                hyperlink_element = adicionar_hyperlink(paragraph, url_limpa, url_limpa)
                paragraph._p.append(hyperlink_element)
                
                # ✅ Adicionar a pontuação como texto normal (não hyperlink)
                if pontuacao:
                    paragraph.add_run(pontuacao)

                # Continuar com o resto do texto
                texto_restante = partes[1]

    # Adicionar texto restante após a última URL (se houver)
    if texto_restante:
        paragraph.add_run(texto_restante)

    return True

def converter_urls_docx_para_hyperlinks(arquivo_entrada, pasta_destino='/app/output', pasta_id_drive=None):
    # 1️⃣ Validar se o arquivo existe
    if not os.path.exists(arquivo_entrada):
        print(f"❌ Erro: Arquivo '{arquivo_entrada}' não encontrado!")
        return False

    print(f"📖 Abrindo arquivo: {arquivo_entrada}")
    doc = Document(arquivo_entrada)

    total_paragrafos_processados = 0
    total_urls_convertidas = 0

    # 2️⃣ Processar parágrafos principais
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            urls_antes = len(re.findall(r'https?://[^\s]+', p.text))
            if processar_urls_em_paragrafo(p):
                total_paragrafos_processados += 1
                total_urls_convertidas += urls_antes
                print(f"   ✅ Parágrafo {i+1} processado com {urls_antes} URLs")

    # 3️⃣ Processar tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        urls_antes = len(re.findall(r'https?://[^\s]+', p.text))
                        if processar_urls_em_paragrafo(p):
                            total_paragrafos_processados += 1
                            total_urls_convertidas += urls_antes

    # 4️⃣ Gerar nome do arquivo final
    nome_base = os.path.basename(arquivo_entrada).replace('.docx', '')
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    arquivo_saida = os.path.join('output', f"{nome_base}_{timestamp}.docx")

    # 5️⃣ Salvar primeiro o arquivo localmente
    os.makedirs('output', exist_ok=True)
    doc.save(arquivo_saida)
    
    print(f"\n📊 Estatísticas do processamento:")
    print(f"   - Parágrafos processados: {total_paragrafos_processados}")
    print(f"   - URLs convertidas em hyperlinks: {total_urls_convertidas}")
    print(f"💾 Arquivo salvo localmente: {arquivo_saida}")

    # 6️⃣ Upload para Google Drive, se configurado
    if pasta_id_drive:
        try:
            upload_para_google_drive(arquivo_saida, os.path.basename(arquivo_saida), pasta_id_drive)
        except Exception as e:
            print(f"⚠️ Erro no upload para Google Drive: {str(e)}")

    # 7️⃣ Copiar para a pasta compartilhada do Docker (opcional)
    if os.path.isdir(pasta_destino):
        destino_drive = os.path.join(pasta_destino, os.path.basename(arquivo_saida))
        try:
            # Evita SameFileError quando origem e destino são o mesmo arquivo
            if not (os.path.exists(destino_drive) and os.path.samefile(arquivo_saida, destino_drive)):
                shutil.copy2(arquivo_saida, destino_drive)
                print(f"📂 Arquivo também copiado para pasta compartilhada: {destino_drive}")
            else:
                print("ℹ️ Origem e destino são o mesmo arquivo; cópia ignorada.")
        except FileNotFoundError:
            # Alguns FS pedem que o diretório exista antes do samefile; garanta e copie
            os.makedirs(pasta_destino, exist_ok=True)
            shutil.copy2(arquivo_saida, destino_drive)
            print(f"📂 Arquivo também copiado para pasta compartilhada: {destino_drive}")
    else:
        print(f"⚠️ Pasta destino '{pasta_destino}' não encontrada. Pulei a cópia local.")

    return True

def metodo_alternativo_melhorado(arquivo_entrada, arquivo_saida):
    """
    Método alternativo melhorado - substitui URLs por texto com formatação
    """
    try:
        print(f"📖 Método alternativo melhorado - Abrindo arquivo: {arquivo_entrada}")
        doc = Document(arquivo_entrada)

        total_urls_encontradas = 0
        paragrafos_processados = 0

        # ✅ Regex corrigida para capturar URLs completas sem caracteres indesejados
        padrao_url = r'https?://[^\s<>"{}|\\^`\[\]\(\),;]+(?:[^\s<>"{}|\\^`\[\]\(\),;.])*'

        # Processar parágrafos
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue

            texto_original = paragraph.text
            urls_no_texto = re.findall(padrao_url, texto_original)
            
            # ✅ Limpeza das URLs encontradas
            urls_limpas = []
            for url in urls_no_texto:
                url_limpa = re.sub(r'[),;.:!?]+$', '', url)
                if url_limpa and url_limpa not in urls_limpas:
                    urls_limpas.append(url_limpa)

            if urls_limpas:
                print(f"   🔗 Parágrafo {paragrafos_processados + 1}: {len(urls_limpas)} URLs encontradas")
                total_urls_encontradas += len(urls_limpas)

                # Limpar o parágrafo
                _limpar_paragrafo(paragraph)

                # Reconstruir o parágrafo com formatação
                texto_restante = texto_original

                for url_limpa in urls_limpas:
                    # Procurar pela URL original no texto
                    padrao_busca = re.escape(url_limpa) + r'[),;.:!?]*'
                    match = re.search(padrao_busca, texto_restante)
                    
                    if match:
                        url_original = match.group()
                        partes = texto_restante.split(url_original, 1)
                        
                        if len(partes) == 2:
                            # Adicionar texto antes da URL
                            if partes[0]:
                                paragraph.add_run(partes[0])

                            # ✅ Usar URL limpa
                            paragraph._p.append(adicionar_hyperlink(paragraph, url_limpa, url_limpa))

                            # Continuar com o resto
                            texto_restante = partes[1]

                # Adicionar texto restante
                if texto_restante:
                    paragraph.add_run(texto_restante)

            paragrafos_processados += 1

        # Processar tabelas também
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if not paragraph.text.strip():
                            continue

                        texto_original = paragraph.text
                        urls_no_texto = re.findall(padrao_url, texto_original)
                        
                        # ✅ Limpeza das URLs encontradas
                        urls_limpas = []
                        for url in urls_no_texto:
                            url_limpa = re.sub(r'[),;.:!?]+$', '', url)
                            if url_limpa and url_limpa not in urls_limpas:
                                urls_limpas.append(url_limpa)

                        if urls_limpas:
                            total_urls_encontradas += len(urls_limpas)
                            _limpar_paragrafo(paragraph)

                            texto_restante = texto_original
                            for url_limpa in urls_limpas:
                                padrao_busca = re.escape(url_limpa) + r'[),;.:!?]*'
                                match = re.search(padrao_busca, texto_restante)
                                
                                if match:
                                    url_original = match.group()
                                    partes = texto_restante.split(url_original, 1)
                                    
                                    if len(partes) == 2:
                                        if partes[0]:
                                            paragraph.add_run(partes[0])

                                        # ✅ Usar URL limpa
                                        paragraph._p.append(adicionar_hyperlink(paragraph, url_limpa, url_limpa))

                                        texto_restante = partes[1]

                            if texto_restante:
                                paragraph.add_run(texto_restante)

        # Salvar documento
        doc.save(arquivo_saida)

        # Upload para Google Drive
        arquivo_local = arquivo_saida
        nome_arquivo = os.path.basename(arquivo_saida)
        try:
            upload_para_google_drive(arquivo_local, nome_arquivo, "1HCo8W9Q9ak8aKOmMRPhSyVBntCS_GD6J")        
        except Exception as e:
            print(f"⚠️ Erro no upload para Google Drive: {str(e)}")

        # Copiar para o Google Drive
        pasta_drive = r'/app/relatorios/'  # Altere para o caminho da sua pasta do Drive
        if os.path.isdir(pasta_drive):
            destino_drive = os.path.join(pasta_drive, os.path.basename(arquivo_saida))
            shutil.copy2(arquivo_saida, destino_drive)
            print(f"📁 Arquivo também salvo em: {destino_drive}")
        else:
            print(f"⚠️ Pasta do Google Drive não encontrada: {pasta_drive}")
            
        print(f"\n✅ Método alternativo concluído!")
        print(f"📊 Estatísticas:")
        print(f"   - Parágrafos processados: {paragrafos_processados}")
        print(f"   - Total de URLs formatadas: {total_urls_encontradas}")
        print(f"💾 Arquivo salvo como: {arquivo_saida}")

        return True

    except Exception as e:
        print(f"❌ Erro no método alternativo: {str(e)}")
        return False

def testar_regex():
    """
    Função para testar a regex com URLs de exemplo
    """
    print("🧪 Testando regex com URLs de exemplo...")

    # URLs de teste
    urls_teste = [
        "https://tinyurl.com/2aymnjlf",
        "https://www.google.com/search?q=python",
        "http://example.com/path/to/file.html",
        "https://github.com/user/repo#readme",
        "https://site.com/page?param1=value1&param2=value2"
    ]

    # Regex melhorada
    padrao_url = r'https?://[^\s<>"{}|\\^`\[\]\(\),;]+(?:[^\s<>"{}|\\^`\[\]\(\),;.])*'

    for url in urls_teste:
        match = re.search(padrao_url, url)
        if match:
            print(f"   ✅ {url} -> Capturado: {match.group()}")
        else:
            print(f"   ❌ {url} -> Não capturado")

    print("\n" + "="*50)

def gerar_versao_ajustada(arquivo_preliminar, pasta_id_drive=None):
    """
    Aplica os ajustes finais ao relatório:
    - Converte URLs em hyperlinks
    - Gera nome do arquivo com timestamp
    - Salva localmente e realiza upload para o Google Drive (se configurado)
    """

    if not os.path.exists(arquivo_preliminar):
        print(f"❌ Arquivo não encontrado: {arquivo_preliminar}")
        return

    print(f"📖 Aplicando versão ajustada com hyperlinks e timestamp...")
    
    # 🧠 Reaproveitar função que já processa os hyperlinks e salva com timestamp
    sucesso = converter_urls_docx_para_hyperlinks(arquivo_preliminar, pasta_id_drive=pasta_id_drive)

    if sucesso:
        print("✅ Versão final ajustada com sucesso.")
    else:
        print("❌ Falha ao gerar a versão ajustada."), '', url)
                pontuacao = url[len(url_limpa):] if len(url) > len(url_limpa) else ''
                
                if url_limpa and url_limpa not in [item[0] for item in urls_processadas]:
                    urls_processadas.append((url_limpa, pontuacao, url))

            if urls_processadas:
                print(f"   🔗 Parágrafo {paragrafos_processados + 1}: {len(urls_processadas)} URLs encontradas")
                total_urls_encontradas += len(urls_processadas)

                # Limpar o parágrafo
                _limpar_paragrafo(paragraph)

                # Reconstruir o parágrafo com formatação
                texto_restante = texto_original

                for url_limpa, pontuacao, url_original in urls_processadas:
                    if url_original in texto_restante:
                        partes = texto_restante.split(url_original, 1)
                        
                        if len(partes) == 2:
                            # Adicionar texto antes da URL
                            if partes[0]:
                                paragraph.add_run(partes[0])

                            # ✅ Usar URL limpa no hyperlink
                            paragraph._p.append(adicionar_hyperlink(paragraph, url_limpa, url_limpa))
                            
                            # ✅ Adicionar pontuação como texto normal
                            if pontuacao:
                                paragraph.add_run(pontuacao)

                            # Continuar com o resto
                            texto_restante = partes[1]

                # Adicionar texto restante
                if texto_restante:
                    paragraph.add_run(texto_restante)

            paragrafos_processados += 1

        # Processar tabelas também
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if not paragraph.text.strip():
                            continue

                        texto_original = paragraph.text
                        urls_no_texto = re.findall(padrao_url, texto_original)
                        
                        # ✅ Limpeza das URLs encontradas
                        urls_limpas = []
                        for url in urls_no_texto:
                            url_limpa = re.sub(r'[),;.:!?]+$', '', url)
                            if url_limpa and url_limpa not in urls_limpas:
                                urls_limpas.append(url_limpa)

                        if urls_limpas:
                            total_urls_encontradas += len(urls_limpas)
                            _limpar_paragrafo(paragraph)

                            texto_restante = texto_original
                            for url_limpa in urls_limpas:
                                padrao_busca = re.escape(url_limpa) + r'[),;.:!?]*'
                                match = re.search(padrao_busca, texto_restante)
                                
                                if match:
                                    url_original = match.group()
                                    partes = texto_restante.split(url_original, 1)
                                    
                                    if len(partes) == 2:
                                        if partes[0]:
                                            paragraph.add_run(partes[0])

                                        # ✅ Usar URL limpa
                                        paragraph._p.append(adicionar_hyperlink(paragraph, url_limpa, url_limpa))

                                        texto_restante = partes[1]

                            if texto_restante:
                                paragraph.add_run(texto_restante)

        # Salvar documento
        doc.save(arquivo_saida)

        # Upload para Google Drive
        arquivo_local = arquivo_saida
        nome_arquivo = os.path.basename(arquivo_saida)
        try:
            upload_para_google_drive(arquivo_local, nome_arquivo, "1HCo8W9Q9ak8aKOmMRPhSyVBntCS_GD6J")        
        except Exception as e:
            print(f"⚠️ Erro no upload para Google Drive: {str(e)}")

        # Copiar para o Google Drive
        pasta_drive = r'/app/relatorios/'  # Altere para o caminho da sua pasta do Drive
        if os.path.isdir(pasta_drive):
            destino_drive = os.path.join(pasta_drive, os.path.basename(arquivo_saida))
            shutil.copy2(arquivo_saida, destino_drive)
            print(f"📁 Arquivo também salvo em: {destino_drive}")
        else:
            print(f"⚠️ Pasta do Google Drive não encontrada: {pasta_drive}")
            
        print(f"\n✅ Método alternativo concluído!")
        print(f"📊 Estatísticas:")
        print(f"   - Parágrafos processados: {paragrafos_processados}")
        print(f"   - Total de URLs formatadas: {total_urls_encontradas}")
        print(f"💾 Arquivo salvo como: {arquivo_saida}")

        return True

    except Exception as e:
        print(f"❌ Erro no método alternativo: {str(e)}")
        return False

def testar_regex():
    """
    Função para testar a regex com URLs de exemplo
    """
    print("🧪 Testando regex com URLs de exemplo...")

    # URLs de teste
    urls_teste = [
        "https://tinyurl.com/2aymnjlf",
        "https://www.google.com/search?q=python",
        "http://example.com/path/to/file.html",
        "https://github.com/user/repo#readme",
        "https://site.com/page?param1=value1&param2=value2"
    ]

    # Regex melhorada
    padrao_url = r'https?://[^\s<>"{}|\\^`\[\]\(\),;]+(?:[^\s<>"{}|\\^`\[\]\(\),;.])*'

    for url in urls_teste:
        match = re.search(padrao_url, url)
        if match:
            print(f"   ✅ {url} -> Capturado: {match.group()}")
        else:
            print(f"   ❌ {url} -> Não capturado")

    print("\n" + "="*50)

def gerar_versao_ajustada(arquivo_preliminar, pasta_id_drive=None):
    """
    Aplica os ajustes finais ao relatório:
    - Converte URLs em hyperlinks
    - Gera nome do arquivo com timestamp
    - Salva localmente e realiza upload para o Google Drive (se configurado)
    """

    if not os.path.exists(arquivo_preliminar):
        print(f"❌ Arquivo não encontrado: {arquivo_preliminar}")
        return

    print(f"📖 Aplicando versão ajustada com hyperlinks e timestamp...")
    
    # 🧠 Reaproveitar função que já processa os hyperlinks e salva com timestamp
    sucesso = converter_urls_docx_para_hyperlinks(arquivo_preliminar, pasta_id_drive=pasta_id_drive)

    if sucesso:
        print("✅ Versão final ajustada com sucesso.")
    else:
        print("❌ Falha ao gerar a versão ajustada."), '', url)
        pontuacao = url[len(url_limpa):] if len(url) > len(url_limpa) else ''
        
        if url_limpa and url_limpa not in [item[0] for item in urls_processadas]:
            urls_processadas.append((url_limpa, pontuacao, url))

    if not urls_processadas:
        return False

    print(f"   🔗 Encontradas {len(urls_processadas)} URLs: {[item[0] for item in urls_processadas[:2]]}{'...' if len(urls_processadas) > 2 else ''}")

    # Limpar o parágrafo atual
    _limpar_paragrafo(paragraph)

    # ✅ PROCESSAMENTO APRIMORADO: Manter pontuação original
    texto_restante = texto_completo

    for url_limpa, pontuacao, url_original in urls_processadas:
        if url_original in texto_restante:
            # Dividir o texto pela URL original
            partes = texto_restante.split(url_original, 1)
            
            if len(partes) == 2:
                # Adicionar texto antes da URL (se houver)
                if partes[0]:
                    paragraph.add_run(partes[0])

                # ✅ Criar hyperlink apenas com URL limpa
                hyperlink_element = adicionar_hyperlink(paragraph, url_limpa, url_limpa)
                paragraph._p.append(hyperlink_element)
                
                # ✅ Adicionar a pontuação como texto normal (não hyperlink)
                if pontuacao:
                    paragraph.add_run(pontuacao)

                # Continuar com o resto do texto
                texto_restante = partes[1]

    # Adicionar texto restante após a última URL (se houver)
    if texto_restante:
        paragraph.add_run(texto_restante)

    return True

def converter_urls_docx_para_hyperlinks(arquivo_entrada, pasta_destino='/app/output', pasta_id_drive=None):
    # 1️⃣ Validar se o arquivo existe
    if not os.path.exists(arquivo_entrada):
        print(f"❌ Erro: Arquivo '{arquivo_entrada}' não encontrado!")
        return False

    print(f"📖 Abrindo arquivo: {arquivo_entrada}")
    doc = Document(arquivo_entrada)

    total_paragrafos_processados = 0
    total_urls_convertidas = 0

    # 2️⃣ Processar parágrafos principais
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            urls_antes = len(re.findall(r'https?://[^\s]+', p.text))
            if processar_urls_em_paragrafo(p):
                total_paragrafos_processados += 1
                total_urls_convertidas += urls_antes
                print(f"   ✅ Parágrafo {i+1} processado com {urls_antes} URLs")

    # 3️⃣ Processar tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        urls_antes = len(re.findall(r'https?://[^\s]+', p.text))
                        if processar_urls_em_paragrafo(p):
                            total_paragrafos_processados += 1
                            total_urls_convertidas += urls_antes

    # 4️⃣ Gerar nome do arquivo final
    nome_base = os.path.basename(arquivo_entrada).replace('.docx', '')
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    arquivo_saida = os.path.join('output', f"{nome_base}_{timestamp}.docx")

    # 5️⃣ Salvar primeiro o arquivo localmente
    os.makedirs('output', exist_ok=True)
    doc.save(arquivo_saida)
    
    print(f"\n📊 Estatísticas do processamento:")
    print(f"   - Parágrafos processados: {total_paragrafos_processados}")
    print(f"   - URLs convertidas em hyperlinks: {total_urls_convertidas}")
    print(f"💾 Arquivo salvo localmente: {arquivo_saida}")

    # 6️⃣ Upload para Google Drive, se configurado
    if pasta_id_drive:
        try:
            upload_para_google_drive(arquivo_saida, os.path.basename(arquivo_saida), pasta_id_drive)
        except Exception as e:
            print(f"⚠️ Erro no upload para Google Drive: {str(e)}")

    # 7️⃣ Copiar para a pasta compartilhada do Docker (opcional)
    if os.path.isdir(pasta_destino):
        destino_drive = os.path.join(pasta_destino, os.path.basename(arquivo_saida))
        try:
            # Evita SameFileError quando origem e destino são o mesmo arquivo
            if not (os.path.exists(destino_drive) and os.path.samefile(arquivo_saida, destino_drive)):
                shutil.copy2(arquivo_saida, destino_drive)
                print(f"📂 Arquivo também copiado para pasta compartilhada: {destino_drive}")
            else:
                print("ℹ️ Origem e destino são o mesmo arquivo; cópia ignorada.")
        except FileNotFoundError:
            # Alguns FS pedem que o diretório exista antes do samefile; garanta e copie
            os.makedirs(pasta_destino, exist_ok=True)
            shutil.copy2(arquivo_saida, destino_drive)
            print(f"📂 Arquivo também copiado para pasta compartilhada: {destino_drive}")
    else:
        print(f"⚠️ Pasta destino '{pasta_destino}' não encontrada. Pulei a cópia local.")

    return True

def metodo_alternativo_melhorado(arquivo_entrada, arquivo_saida):
    """
    Método alternativo melhorado - substitui URLs por texto com formatação
    """
    try:
        print(f"📖 Método alternativo melhorado - Abrindo arquivo: {arquivo_entrada}")
        doc = Document(arquivo_entrada)

        total_urls_encontradas = 0
        paragrafos_processados = 0

        # ✅ Regex corrigida para capturar URLs completas sem caracteres indesejados
        padrao_url = r'https?://[^\s<>"{}|\\^`\[\]\(\),;]+(?:[^\s<>"{}|\\^`\[\]\(\),;.])*'

        # Processar parágrafos
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue

            texto_original = paragraph.text
            urls_no_texto = re.findall(padrao_url, texto_original)
            
            # ✅ Limpeza das URLs encontradas
            urls_limpas = []
            for url in urls_no_texto:
                url_limpa = re.sub(r'[),;.:!?]+$', '', url)
                if url_limpa and url_limpa not in urls_limpas:
                    urls_limpas.append(url_limpa)

            if urls_limpas:
                print(f"   🔗 Parágrafo {paragrafos_processados + 1}: {len(urls_limpas)} URLs encontradas")
                total_urls_encontradas += len(urls_limpas)

                # Limpar o parágrafo
                _limpar_paragrafo(paragraph)

                # Reconstruir o parágrafo com formatação
                texto_restante = texto_original

                for url_limpa in urls_limpas:
                    # Procurar pela URL original no texto
                    padrao_busca = re.escape(url_limpa) + r'[),;.:!?]*'
                    match = re.search(padrao_busca, texto_restante)
                    
                    if match:
                        url_original = match.group()
                        partes = texto_restante.split(url_original, 1)
                        
                        if len(partes) == 2:
                            # Adicionar texto antes da URL
                            if partes[0]:
                                paragraph.add_run(partes[0])

                            # ✅ Usar URL limpa
                            paragraph._p.append(adicionar_hyperlink(paragraph, url_limpa, url_limpa))

                            # Continuar com o resto
                            texto_restante = partes[1]

                # Adicionar texto restante
                if texto_restante:
                    paragraph.add_run(texto_restante)

            paragrafos_processados += 1

        # Processar tabelas também
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if not paragraph.text.strip():
                            continue

                        texto_original = paragraph.text
                        urls_no_texto = re.findall(padrao_url, texto_original)
                        
                        # ✅ Limpeza das URLs encontradas
                        urls_limpas = []
                        for url in urls_no_texto:
                            url_limpa = re.sub(r'[),;.:!?]+$', '', url)
                            if url_limpa and url_limpa not in urls_limpas:
                                urls_limpas.append(url_limpa)

                        if urls_limpas:
                            total_urls_encontradas += len(urls_limpas)
                            _limpar_paragrafo(paragraph)

                            texto_restante = texto_original
                            for url_limpa in urls_limpas:
                                padrao_busca = re.escape(url_limpa) + r'[),;.:!?]*'
                                match = re.search(padrao_busca, texto_restante)
                                
                                if match:
                                    url_original = match.group()
                                    partes = texto_restante.split(url_original, 1)
                                    
                                    if len(partes) == 2:
                                        if partes[0]:
                                            paragraph.add_run(partes[0])

                                        # ✅ Usar URL limpa
                                        paragraph._p.append(adicionar_hyperlink(paragraph, url_limpa, url_limpa))

                                        texto_restante = partes[1]

                            if texto_restante:
                                paragraph.add_run(texto_restante)

        # Salvar documento
        doc.save(arquivo_saida)

        # Upload para Google Drive
        arquivo_local = arquivo_saida
        nome_arquivo = os.path.basename(arquivo_saida)
        try:
            upload_para_google_drive(arquivo_local, nome_arquivo, "1HCo8W9Q9ak8aKOmMRPhSyVBntCS_GD6J")        
        except Exception as e:
            print(f"⚠️ Erro no upload para Google Drive: {str(e)}")

        # Copiar para o Google Drive
        pasta_drive = r'/app/relatorios/'  # Altere para o caminho da sua pasta do Drive
        if os.path.isdir(pasta_drive):
            destino_drive = os.path.join(pasta_drive, os.path.basename(arquivo_saida))
            shutil.copy2(arquivo_saida, destino_drive)
            print(f"📁 Arquivo também salvo em: {destino_drive}")
        else:
            print(f"⚠️ Pasta do Google Drive não encontrada: {pasta_drive}")
            
        print(f"\n✅ Método alternativo concluído!")
        print(f"📊 Estatísticas:")
        print(f"   - Parágrafos processados: {paragrafos_processados}")
        print(f"   - Total de URLs formatadas: {total_urls_encontradas}")
        print(f"💾 Arquivo salvo como: {arquivo_saida}")

        return True

    except Exception as e:
        print(f"❌ Erro no método alternativo: {str(e)}")
        return False

def testar_regex():
    """
    Função para testar a regex com URLs de exemplo
    """
    print("🧪 Testando regex com URLs de exemplo...")

    # URLs de teste
    urls_teste = [
        "https://tinyurl.com/2aymnjlf",
        "https://www.google.com/search?q=python",
        "http://example.com/path/to/file.html",
        "https://github.com/user/repo#readme",
        "https://site.com/page?param1=value1&param2=value2"
    ]

    # Regex melhorada
    padrao_url = r'https?://[^\s<>"{}|\\^`\[\]\(\),;]+(?:[^\s<>"{}|\\^`\[\]\(\),;.])*'

    for url in urls_teste:
        match = re.search(padrao_url, url)
        if match:
            print(f"   ✅ {url} -> Capturado: {match.group()}")
        else:
            print(f"   ❌ {url} -> Não capturado")

    print("\n" + "="*50)

def gerar_versao_ajustada(arquivo_preliminar, pasta_id_drive=None):
    """
    Aplica os ajustes finais ao relatório:
    - Converte URLs em hyperlinks
    - Gera nome do arquivo com timestamp
    - Salva localmente e realiza upload para o Google Drive (se configurado)
    """

    if not os.path.exists(arquivo_preliminar):
        print(f"❌ Arquivo não encontrado: {arquivo_preliminar}")
        return

    print(f"📖 Aplicando versão ajustada com hyperlinks e timestamp...")
    
    # 🧠 Reaproveitar função que já processa os hyperlinks e salva com timestamp
    sucesso = converter_urls_docx_para_hyperlinks(arquivo_preliminar, pasta_id_drive=pasta_id_drive)

    if sucesso:
        print("✅ Versão final ajustada com sucesso.")
    else:
        print("❌ Falha ao gerar a versão ajustada.")