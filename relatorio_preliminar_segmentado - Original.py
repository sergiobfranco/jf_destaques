# Versão Segmentada do Relatório

import pandas as pd
import pyshorteners
from docx import Document
from docx.shared import Pt  # Import Pt for point size
from docx.enum.text import WD_LINE_SPACING  # Import line spacing enum
import re
import time # Importar time para pausa
from config import arq_resumo_final, marca1, marca2
import requests
import os
from datetime import datetime

# 1. FUNÇÃO AUXILIAR PARA ENCURTAMENTO DE URL (adicionar no início do arquivo)
def encurtar_url_seguro(url_original, max_tentativas=3, delay=2):
    """
    Função auxiliar para encurtar URLs de forma segura com tratamento de erro robusto
    """
    import pyshorteners
    import requests
    import time
    
    # Validar se a URL é válida antes de tentar encurtar
    if not url_original or pd.isna(url_original) or str(url_original).strip() == '':
        print(f"URL vazia ou inválida: {url_original}")
        return str(url_original) if url_original else "URL não disponível"
    
    url_str = str(url_original).strip()
    
    # Verificar se já é uma URL válida
    if not url_str.startswith(('http://', 'https://')):
        print(f"URL não tem protocolo válido: {url_str}")
        return url_str
    
    s = pyshorteners.Shortener()
    
    for tentativa in range(max_tentativas):
        try:
            # Tentar TinyURL primeiro
            short_url = s.tinyurl.short(url_str)
            return short_url
            
        except Exception as e:
            erro_str = str(e)
            print(f"Tentativa {tentativa + 1}/{max_tentativas} - Erro TinyURL: {erro_str}")
            
            # Se for erro específico do TinyURL, tentar serviço alternativo
            if tentativa == max_tentativas - 1:
                try:
                    # Tentar serviço alternativo (is.gd)
                    short_url = s.isgd.short(url_str)
                    print(f"URL encurtada com serviço alternativo (is.gd)")
                    return short_url
                except Exception as e2:
                    print(f"Erro também no serviço alternativo: {str(e2)}")
                    break
            
            if tentativa < max_tentativas - 1:
                time.sleep(delay)
    
    print(f"Falha em todos os serviços de encurtamento. Usando URL original.")
    return url_str

def processar_editoriais_integrado(final_df_editorial, document, opcao_selecionada=1):
    """
    Processa editoriais de forma integrada (substitui toda a seção de editoriais)
    """
    print(f"\nProcessando {len(final_df_editorial)} editoriais...")
    
    if final_df_editorial.empty:
        print("Nenhum editorial encontrado.")
        return

    # Debug: verificar estrutura do DataFrame
    print(f"Colunas disponíveis em final_df_editorial: {list(final_df_editorial.columns)}")
    print(f"Primeiras linhas do DataFrame:")
    print(final_df_editorial.head())

    for index, row_editorial in final_df_editorial.iterrows():
        try:
            # Tentar diferentes nomes de colunas possíveis
            w_veiculo_editorial = (
                row_editorial.get('Veiculo') or 
                row_editorial.get('Veículo') or 
                row_editorial.get('veiculo') or 
                'Veículo Desconhecido'
            )
            
            w_titulo_editorial = (
                row_editorial.get('Titulo') or 
                row_editorial.get('Título') or 
                row_editorial.get('titulo') or 
                row_editorial.get('Conteudo') or 
                row_editorial.get('Conteúdo') or 
                'Título não disponível'
            )
            
            w_url_editorial = (
                row_editorial.get('UrlVisualizacao') or 
                row_editorial.get('Url') or 
                row_editorial.get('URL') or 
                row_editorial.get('Link') or 
                'URL Não Disponível'
            )
            
            print(f"\nEditorial {index + 1}:")
            print(f"  Veículo: {w_veiculo_editorial}")
            print(f"  Título: {w_titulo_editorial}")
            print(f"  URL: {w_url_editorial}")
            
            # Adicionar informações básicas
            document.add_paragraph(f"{w_veiculo_editorial}: {w_titulo_editorial}")
            
            # Encurtamento seguro da URL
            if w_url_editorial and w_url_editorial != 'URL Não Disponível':
                short_url_editorial = encurtar_url_seguro(w_url_editorial, max_tentativas=3, delay=1)
            else:
                short_url_editorial = w_url_editorial
                print(f"URL inválida para editorial {index + 1}, usando valor original")
            
            document.add_paragraph(short_url_editorial)
            document.add_paragraph("*")
            
            print(f"Editorial {index + 1} processado com sucesso")
            
        except Exception as e:
            print(f"Erro ao processar editorial {index + 1}: {str(e)}")
            print(f"Dados do editorial: {dict(row_editorial)}")
            # Continuar com o próximo editorial em caso de erro
            document.add_paragraph(f"Editorial {index + 1}: Erro no processamento")
            document.add_paragraph("URL não disponível")
            document.add_paragraph("*")
            continue

def debug_dataframe_structure(df, nome_df):
    """
    Função auxiliar para debugar a estrutura dos DataFrames
    """
    print(f"\n=== DEBUG: Estrutura do {nome_df} ===")
    print(f"Shape: {df.shape}")
    print(f"Colunas: {list(df.columns)}")
    if not df.empty:
        print(f"Primeiras 2 linhas:")
        for i in range(min(2, len(df))):
            print(f"Linha {i}: {dict(df.iloc[i])}")
    else:
        print("DataFrame vazio")
    print("=" * 50)

def obter_nome_veiculo(codigo_veiculo):
    """
    Retorna o nome do veículo baseado no código
    """
    nomes_veiculos = {
        675: "O ESTADO DE S.PAULO/SÃO PAULO",
        10459: "VALOR ECONÔMICO / SÃO PAULO", 
        331: "FOLHA DE S.PAULO / SÃO PAULO",
        682: "O GLOBO / RIO DE JANEIRO"
    }
    return nomes_veiculos.get(codigo_veiculo, f"VEÍCULO CÓDIGO {codigo_veiculo}")

def gerar_versao_preliminar(final_df_small_marca, final_df_small_marca_irrelevantes, df_resumo_marca, 
                          df_resumo_marca_irrelevantes, final_df_marca, df_resumo_setor, final_df_setor, 
                          final_df_editorial, final_df_SPECIALS_small, opcao_selecionada=1, codigo_veiculo=None):
    """
    Gera versão preliminar do relatório baseada na opção selecionada
    
    Args:
        opcao_selecionada (int): Opção selecionada (1-7)
        codigo_veiculo (int): Código do veículo (para opções 4-7)
    """
    
    # Determinar quais seções executar baseado na opção selecionada
    executar_marcas = opcao_selecionada in [1, 2]  # Relatório Completo ou Somente Marcas
    executar_setor = opcao_selecionada in [1, 3, 4, 5, 6, 7]  # Todas exceto Somente Marcas
    executar_editoriais = opcao_selecionada in [1, 2, 3]  # Não executar para opções específicas de veículos
    
    print(f"\n=== CONFIGURAÇÃO DE PROCESSAMENTO ===")
    print(f"Opção selecionada: {opcao_selecionada}")
    print(f"Código do veículo: {codigo_veiculo}")
    print(f"Executar MARCAS: {executar_marcas}")
    print(f"Executar SETOR: {executar_setor}")
    print(f"Executar EDITORIAIS: {executar_editoriais}")
    print("=" * 50)
    
    # 0. SEÇÃO DE CONFIGURAÇÕES INICIAIS (SEMPRE EXECUTADA)
    print("\n=== SEÇÃO 0: CONFIGURAÇÕES INICIAIS ===")
    
    # DEBUG: Verificar estrutura dos DataFrames recebidos
    debug_dataframe_structure(final_df_editorial, "final_df_editorial")

    # Inicializar o documento DOCX
    document = Document()

    # Configurar o estilo Normal para remover espaço após o parágrafo e usar espaçamento simples
    styles = document.styles
    style = styles['Normal']
    font = style.font
    font.name = 'Calibri'  # Set the font to Calibri
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Adicionar título e linha em branco
    document.add_paragraph("DESTAQUES DO DIA J&F")
    document.add_paragraph("")
    
    # 1. SEÇÃO DE MARCAS (CONDICIONAL)
    if executar_marcas:
        print("\n=== SEÇÃO 1: MARCAS ===")
        
        document.add_paragraph("--- NOTÍCIAS DE MARCAS ---")
        document.add_paragraph("")

        # --- Seção Original para resumos de Marca ---
        print(f"Processando {len(df_resumo_marca)} resumos de Marca...")
        for index, row_marca in df_resumo_marca.iterrows():
            # Inicializar a string para cada grupo de notícias
            group_string = ""

            # Iterar sobre os IDs das notícias (Marca)
            if 'Ids' not in row_marca or pd.isna(row_marca['Ids']):
                print(f"Aviso: Linha {index} no df_resumo_marca não tem IDs válidos. Pulando.")
                continue

            for news_id_str in str(row_marca['Ids']).split(','):
                try:
                    news_id = int(news_id_str.strip())
                except ValueError:
                    print(f"Aviso: Não foi possível converter ID '{news_id_str}' para inteiro na linha {index} de df_resumo_marca. Pulando.")
                    continue

                # Consultar informações no final_df_marca
                news_info_marca = final_df_marca[final_df_marca['Id'] == news_id]
                if news_info_marca.empty:
                    print(f"Aviso: ID {news_id} não encontrado em final_df_marca para resumo de Marca. Pulando.")
                    continue

                news_info_marca = news_info_marca.iloc[0]
                w_veiculo_marca = news_info_marca['Veiculo']
                w_url_marca = news_info_marca['UrlVisualizacao']

                # Encurtar a URL
                short_url_marca = encurtar_url_seguro(w_url_marca, max_tentativas=3, delay=1)

                # Atualizar o DataFrame
                if news_id in final_df_small_marca['Id'].values:
                    final_df_small_marca.loc[final_df_small_marca['Id'] == news_id, 'ShortURL'] = short_url_marca
                else:
                    print(f"Aviso: ID {news_id} não encontrado em final_df_small_marca para adicionar ShortURL.")

                # Identificar o tipo "Special"
                special_type = ""
                special_info = final_df_SPECIALS_small[final_df_SPECIALS_small['Id'] == news_id]

                if not special_info.empty:
                    canais_special = special_info.iloc[0]['Canais']
                    if isinstance(canais_special, list):
                        canais_str = ', '.join(map(str, canais_special))
                    else:
                        canais_str = str(canais_special)

                    if "Editoriais" in canais_str:
                        special_type = "Editorial"
                    elif "Colunistas" in canais_str:
                        special_type = "Colunista"
                    elif "1ª Página" in canais_str:
                        special_type = "Capa"

                # Formatar a string
                if special_type:
                    group_string += f"{w_veiculo_marca} ({special_type} - {short_url_marca}), "
                else:
                    group_string += f"{w_veiculo_marca} ({short_url_marca}), "

            # Limpar e adicionar o resumo à string
            group_string = group_string.rstrip(', ')
            group_string += "\n"

            # Adicionar resumo
            if 'Resumo' not in row_marca or pd.isna(row_marca['Resumo']):
                print(f"Aviso: Linha {index} no df_resumo_marca não tem Resumo. Adicionando placeholder.")
                resumo_limpo = "[Resumo não disponível]"
            else:
                resumo_limpo = str(row_marca['Resumo'])

            group_string += resumo_limpo

            # Adicionar a string ao documento DOCX
            document.add_paragraph(group_string)
            document.add_paragraph("")

        # --- Seção para resumos de Marca - NOTÍCIAS IRRELEVANTES ---
        document.add_paragraph("")
        document.add_paragraph("--- CITAÇÕES ---")
        document.add_paragraph("")
        
        print(f"Processando {len(df_resumo_marca_irrelevantes)} resumos de Marca - CITAÇÕES...")
        for index, row_marca in df_resumo_marca_irrelevantes.iterrows():
            group_string = ""

            if 'Ids' not in row_marca or pd.isna(row_marca['Ids']):
                print(f"Aviso: Linha {index} no df_resumo_marca_irrelevantes não tem IDs válidos. Pulando.")
                continue

            for news_id_str in str(row_marca['Ids']).split(','):
                try:
                    news_id = int(news_id_str.strip())
                except ValueError:
                    print(f"Aviso: Não foi possível converter ID '{news_id_str}' para inteiro na linha {index} de df_resumo_marca. Pulando.")
                    continue

                # Consultar informações no final_df_marca
                news_info_marca = final_df_marca[final_df_marca['Id'] == news_id]
                if news_info_marca.empty:
                    print(f"Aviso: ID {news_id} não encontrado em final_df_marca para resumo de Marca. Pulando.")
                    continue

                news_info_marca = news_info_marca.iloc[0]
                w_veiculo_marca = news_info_marca['Veiculo']
                w_url_marca = news_info_marca['UrlVisualizacao']

                # Encurtar a URL
                short_url_marca = encurtar_url_seguro(w_url_marca, max_tentativas=3, delay=1)

                # Atualizar o DataFrame
                if news_id in final_df_small_marca_irrelevantes['Id'].values:
                    final_df_small_marca_irrelevantes.loc[final_df_small_marca_irrelevantes['Id'] == news_id, 'ShortURL'] = short_url_marca

                # Identificar o tipo "Special"
                special_type = ""
                special_info = final_df_SPECIALS_small[final_df_SPECIALS_small['Id'] == news_id]

                if not special_info.empty:
                    canais_special = special_info.iloc[0]['Canais']
                    if isinstance(canais_special, list):
                        canais_str = ', '.join(map(str, canais_special))
                    else:
                        canais_str = str(canais_special)

                    if "Editoriais" in canais_str:
                        special_type = "Editorial"
                    elif "Colunistas" in canais_str:
                        special_type = "Colunista"
                    elif "1ª Página" in canais_str:
                        special_type = "Capa"

                # Formatar a string
                if special_type:
                    group_string += f"{w_veiculo_marca} ({special_type} - {short_url_marca}), "
                else:
                    group_string += f"{w_veiculo_marca} ({short_url_marca}), "

            # Limpar e adicionar o resumo à string
            group_string = group_string.rstrip(', ')
            group_string += "\n"

            # Adicionar resumo
            if 'Resumo' not in row_marca or pd.isna(row_marca['Resumo']):
                print(f"Aviso: Linha {index} no df_resumo_marca_irrelevantes não tem Resumo. Adicionando placeholder.")
                resumo_limpo = "[Resumo não disponível]"
            else:
                resumo_limpo = str(row_marca['Resumo'])

            group_string += resumo_limpo

            # Adicionar a string ao documento DOCX
            document.add_paragraph(group_string)
            document.add_paragraph("")

        # --- Seção de links das notícias por Marcas ---
        final_df_small_marca_combined = pd.concat([final_df_small_marca, final_df_small_marca_irrelevantes], ignore_index=True)

        print(f"\nVerificando final_df_small_marca_combined antes de ordenar para links:")
        print(f"  Tem {len(final_df_small_marca_combined)} linhas.")
        print(f"  Tem coluna 'Canais'? {'Canais' in final_df_small_marca_combined.columns}")

        if not final_df_small_marca_combined.empty and 'Canais' in final_df_small_marca_combined.columns:
            try:
                order = [marca1, marca2] + [marca for marca in final_df_small_marca_combined['Canais'].unique() if marca not in (marca1, marca2)]
                final_df_small_marca_sorted = final_df_small_marca_combined.sort_values(by=['Canais'], key=lambda x: pd.Categorical(x, categories=order, ordered=True))
                print("Ordenação personalizada dos links por Marca aplicada.")
            except (NameError, KeyError, AttributeError) as e:
                print(f"Aviso: Falha na ordenação personalizada dos links por Marca ({type(e).__name__}: {e}).")
                print("Ordenando por Canais padrão.")
                final_df_small_marca_sorted = final_df_small_marca_combined.sort_values(by=['Canais'])

            document.add_paragraph("")
            document.add_paragraph("--- LINKS DAS NOTÍCIAS DE MARCA ---")
            document.add_paragraph("")

            current_marca = None

            for index, row_small_marca in final_df_small_marca_sorted.iterrows():
                marca = row_small_marca['Canais']
                if marca != current_marca:
                    if current_marca is not None:
                        document.add_paragraph("")
                    document.add_paragraph(f"*{marca}*")
                    current_marca = marca

                news_id_small = row_small_marca['Id']
                original_row_marca = final_df_marca[final_df_marca['Id'] == news_id_small]
                if original_row_marca.empty:
                    print(f"Aviso: ID {news_id_small} não encontrado em final_df_marca para links de Marca. Pulando este link.")
                    continue

                original_row_marca = original_row_marca.iloc[0]
                veiculo = original_row_marca['Veiculo']
                titulo = original_row_marca['Titulo']
                canais_commodities = original_row_marca.get('CanaisCommodities', '')

                document.add_paragraph(f"{veiculo}: {titulo}")

                short_url = row_small_marca.get('ShortURL')
                if pd.isna(short_url) or not short_url:
                    short_url = original_row_marca.get('UrlVisualizacao', 'URL Não Encontrada')

                prefix = "Coluna - " if "Colunistas" in str(canais_commodities) else ""
                document.add_paragraph(f"{prefix}{short_url}")
                document.add_paragraph("*")
        else:
            print("DataFrame 'final_df_small_marca_combined' não encontrado, vazio ou sem a coluna 'Canais'. Pulando a seção de links por Marcas.")

    # 2. SEÇÃO DE SETOR (CONDICIONAL)
    if executar_setor:
        print("\n=== SEÇÃO 2: SETOR ===")
        
        print(f"\nVerificando df_resumo_setor:")
        print(f"  Tem {len(df_resumo_setor)} linhas.")
        print(f"  Está vazio? {df_resumo_setor.empty}")
        
        if not df_resumo_setor.empty:
            document.add_paragraph("")
            document.add_paragraph("--- NOTÍCIAS DE SETOR ---")
            
            # Adicionar título específico do veículo se for opção 4-7
            if codigo_veiculo:
                nome_veiculo = obter_nome_veiculo(codigo_veiculo)
                document.add_paragraph(nome_veiculo)
                print(f"Adicionado título específico do veículo: {nome_veiculo}")
            
            document.add_paragraph("")

            # Iterar o dataframe df_resumo_setor
            current_tema = None

            for index, row_setor in df_resumo_setor.iterrows():
                # Check if 'Tema' column exists and is not None
                if 'Tema' not in row_setor or pd.isna(row_setor['Tema']):
                    print(f"Aviso: Linha {index} no df_resumo_setor não tem Tema. Pulando.")
                    continue

                tema = row_setor['Tema']

                # Quando iniciar um Tema novo, incluir uma linha com o conteúdo do campo Tema
                if tema != current_tema:
                    if current_tema is not None:
                        document.add_paragraph("")
                    document.add_paragraph(f"*{tema}*")
                    current_tema = tema

                # Buscar informações no final_df_setor usando o Id
                if 'Id' not in row_setor or pd.isna(row_setor['Id']):
                    print(f"Aviso: Linha {index} no df_resumo_setor não tem Id válido. Pulando.")
                    continue

                try:
                    news_id = int(str(row_setor['Id']).strip())
                except ValueError:
                    print(f"Aviso: Não foi possível converter ID '{row_setor['Id']}' para inteiro na linha {index} de df_resumo_setor. Pulando.")
                    continue

                news_info_setor = final_df_setor[final_df_setor['Id'] == news_id]
                if news_info_setor.empty:
                    print(f"Aviso: ID {news_id} não encontrado em final_df_setor para resumo de Setor. Pulando.")
                    continue
                news_info_setor = news_info_setor.iloc[0]

                w_veiculo_setor = news_info_setor['Veiculo']
                w_titulo_setor = news_info_setor['Titulo']
                w_url_setor = news_info_setor['UrlVisualizacao']

                # Incluir linha com Veiculo e Titulo
                document.add_paragraph(f"{w_veiculo_setor}: {w_titulo_setor}")

                # Incluir o Resumo do Setor
                if 'Resumo' not in row_setor or pd.isna(row_setor['Resumo']):
                    print(f"Aviso: Linha {index} no df_resumo_setor (Tema {tema}) não tem Resumo. Adicionando placeholder.")
                    document.add_paragraph("[Resumo não disponível]")
                else:
                    resumo_bruto = str(row_setor['Resumo'])
                    document.add_paragraph(resumo_bruto)

                # Incluir a URL encurtada
                short_url_setor = encurtar_url_seguro(w_url_setor, max_tentativas=3, delay=1)
                document.add_paragraph(short_url_setor)

                # Incluir a linha com um asterisco
                document.add_paragraph("*")

    # 3. SEÇÃO DE EDITORIAIS (CONDICIONAL)
    if executar_editoriais:
        print("\n=== SEÇÃO 3: EDITORIAIS ===")
        
        # Adicionar cabeçalho da seção de editoriais
        document.add_paragraph("")
        document.add_paragraph("--- EDITORIAIS ---")    
        document.add_paragraph("")

        # Processar Editoriais
        processar_editoriais_integrado(final_df_editorial, document, opcao_selecionada)

    # 4. SEÇÃO DE FINALIZAÇÃO (SEMPRE EXECUTADA)
    print("\n=== SEÇÃO 4: FINALIZAÇÃO ===")
    
    # ===== PROCESSAMENTO DO DOCUMENTO ANTES DA GRAVAÇÃO =====
    print("\nProcessando documento antes da gravação para remover padrões...")

    paragraphs_to_remove_indices = []

    # Pattern for lines starting with "*(" and ending with "palavras)"
    pattern_line_start_paren_end_palavras = r"^\s*\*\s*\(.*?palavras\)\s*$"
    # Pattern for lines starting with "*Resumo"
    pattern_line_start_resumo = r"^\s*\*Resumo.*$"
    # Pattern for any occurrence of "(*...palavras*)" or "(...palavras)" or "**Resumo (...):**" etc.
    pattern_parenthesized_palavras_general = r"\s*[\*\s]*\(.*?\s*palavras\s*\)[\s\:\*]*"
    # Add patterns for specific prefixes like "**Resumo:**" or "*Resumo:*"
    pattern_specific_resumo_prefixes = r"^\s*[\*\s]*Resumo\s*[:\*\s]*"

    for i, paragraph in enumerate(document.paragraphs):
        texto = paragraph.text

        # 1) Eliminar todas as linhas que começam com a string "*(" E terminam com a string "palavras)"
        if re.fullmatch(pattern_line_start_paren_end_palavras, texto, re.IGNORECASE):
            paragraphs_to_remove_indices.append(i)
            print(f"  Removendo linha {i} por corresponder ao padrão '* (... palavras)'")
            continue

        # 2) Eliminar as linhas que começam com a string "*Resumo"
        if re.fullmatch(pattern_line_start_resumo, texto, re.IGNORECASE):
            paragraphs_to_remove_indices.append(i)
            print(f"  Removendo linha {i} por começar com '*Resumo'")
            continue

        # 3) Eliminar as strings que começam com a string "*(", e que tenham algum conteúdo na sequência, e em seguida encerrem com a string "palavras)"
        new_text = re.sub(pattern_parenthesized_palavras_general, '', texto, flags=re.IGNORECASE).strip()
        new_text = re.sub(pattern_specific_resumo_prefixes, '', new_text, flags=re.IGNORECASE).strip()

        # If the text changed, update the paragraph
        if new_text != texto.strip():
            paragraph.text = new_text
            print(f"  Removido padrão '(... palavras)' ou prefixos de resumo na linha {i}")

        # Replace any remaining "**" with "*"
        if "**" in paragraph.text:
            paragraph.text = paragraph.text.replace("**", "*")

    # Remover os parágrafos marcados para remoção (em ordem reversa)
    for i in sorted(paragraphs_to_remove_indices, reverse=True):
        p = document.paragraphs[i]._element
        p.getparent().remove(p)

    print(f"\nTotal de parágrafos removidos: {len(paragraphs_to_remove_indices)}")

    # Salvar o documento processado
    document.save(arq_resumo_final)
    print(f"Arquivo DOCX salvo em: {arq_resumo_final}")
    
    print(f"\n=== RELATÓRIO GERADO COM SUCESSO ===")
    print(f"Opção processada: {opcao_selecionada}")
    if codigo_veiculo:
        print(f"Veículo: {obter_nome_veiculo(codigo_veiculo)}")
    print("=" * 50)