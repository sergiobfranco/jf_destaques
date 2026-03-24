# verificador_marcas_chk.py
# Módulo de verificação de marcas em relatórios DOCX
# Gera arquivo _CHK com manchetes e frases de ação por marca, usando DeepSeek
# Versão 1.0 — Março/2026

import os
import re
import json
import time
import requests
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# ==============================================================================
# CONFIGURAÇÕES
# ==============================================================================

MARCAS_MONITORADAS = [
    'J&F Mineração/LHG Mining',  # Nomes mais longos primeiro (evita match parcial)
    'J&F Mineração/LHG',
    'Joesley Batista',
    'Wesley Batista',
    'Júnior Friboi',
    'Banco Original',
    'Âmbar Energia',
    'Ambar Energia',
    'Instituto J&F',
    'Canal Rural',
    'Braskem',
    'Eldorado',
    'PicPay',
    'Holding',
    'Flora',
    'Ambar',
    'J&F',
    'JBS',
]

SECOES_RELEVANTES = ['--- NOTÍCIAS DE MARCAS ---', '--- CITAÇÕES ---']

# Número de linhas não-vazias a verificar para encontrar a seção de marcas
LINHAS_VERIFICACAO = 30

# Temperatura DeepSeek para geração de manchetes/ações (baixa = mais determinístico)
TEMPERATURA_DEEPSEEK = 0.3

# Máximo de tokens na resposta DeepSeek por parágrafo
MAX_TOKENS_PARAGRAFO = 900

# Pausa entre chamadas à API (segundos) para evitar rate limit
PAUSA_ENTRE_CHAMADAS = 0.5


# ==============================================================================
# UTILITÁRIOS DE API
# ==============================================================================

def _obter_chave_deepseek():
    """
    Obtém a chave da API DeepSeek via .env + config_usuario.ini,
    usando o mesmo padrão do resumos_marcas_v2.py.
    """
    try:
        from dotenv import load_dotenv
        import configparser

        load_dotenv()
        config_path = os.path.join("dados", "config", "config_usuario.ini")
        config = configparser.ConfigParser()
        config.read(config_path, encoding="utf-8")
        perfil = config.get("usuario", "perfil", fallback="client").strip().lower()

        env_var = f"DEEPSEEK_API_KEY_{perfil.upper()}"
        chave = os.getenv(env_var)

        if not chave:
            # Fallback: tentar chave genérica
            chave = os.getenv("DEEPSEEK_API_KEY")

        if not chave:
            raise ValueError(
                f"Chave DeepSeek não encontrada. "
                f"Variável esperada: '{env_var}' (perfil='{perfil}'). "
                f"Também tentou 'DEEPSEEK_API_KEY' sem sucesso."
            )

        return chave

    except ImportError:
        # Se dotenv/configparser não disponíveis, tenta variável de ambiente direta
        chave = os.getenv("DEEPSEEK_API_KEY")
        if not chave:
            raise ValueError("DEEPSEEK_API_KEY não encontrada no ambiente.")
        return chave


def _obter_url_deepseek():
    """Obtém a URL da API DeepSeek (tenta importar de config, fallback hardcoded)."""
    try:
        from config import DEEPSEEK_API_URL
        return DEEPSEEK_API_URL
    except ImportError:
        return "https://api.deepseek.com/v1/chat/completions"


def _chamar_deepseek(prompt, chave_api, url_api, temperatura=0.3, max_tokens=900):
    """
    Faz uma chamada à API DeepSeek e retorna o texto da resposta.
    Lança exceção em caso de falha.
    """
    headers = {
        "Authorization": f"Bearer {chave_api}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": temperatura,
        "max_tokens": max_tokens
    }

    resp = requests.post(url_api, headers=headers, json=payload, timeout=90)
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"].strip()


# ==============================================================================
# DETECÇÃO DE MARCAS
# ==============================================================================

def detectar_marcas_no_paragrafo(texto):
    """
    Retorna lista de marcas (sem repetição) encontradas no texto do parágrafo.
    Remove asteriscos antes de comparar (o relatório usa *marca* para negrito).
    Preserva ordem de aparição no texto.
    """
    texto_limpo = texto.replace('*', '')
    marcas_encontradas = []

    for marca in MARCAS_MONITORADAS:
        # Busca case-insensitive, palavra inteira (ou parte de expressão composta)
        padrao = re.compile(re.escape(marca), re.IGNORECASE)
        if padrao.search(texto_limpo) and marca not in marcas_encontradas:
            marcas_encontradas.append(marca)

    return marcas_encontradas


# ==============================================================================
# GERAÇÃO DE MANCHETE E AÇÕES VIA DEEPSEEK
# ==============================================================================

def gerar_manchete_e_acoes(paragrafo_texto, marcas_encontradas, chave_api, url_api):
    """
    Chama DeepSeek para gerar:
      1. Uma manchete de jornal resumindo o parágrafo
      2. Para cada marca no parágrafo, frases objetivas descrevendo a ação da marca

    Retorna (manchete: str, acoes: list[str])
    """
    marcas_str = ', '.join(marcas_encontradas)

    # Limpar texto do parágrafo (remover asteriscos para o prompt)
    texto_para_prompt = paragrafo_texto.replace('*', '')

    prompt = f"""Analise o parágrafo abaixo e retorne SOMENTE um objeto JSON válido, sem markdown, sem blocos de código, sem nenhum texto adicional antes ou depois.

O JSON deve ter EXATAMENTE este formato:
{{
  "manchete": "Uma manchete curta e objetiva (máx. 15 palavras) que capture o tema central do parágrafo",
  "acoes": [
    "Frase 1: ação da marca X",
    "Frase 2: ação da marca Y",
    "Frase 3: outra ação da marca X (se houver ação diferente)"
  ]
}}

REGRAS PARA A MANCHETE:
- Seja factual, direto, no estilo manchete de jornal
- Máximo 15 palavras
- Não use aspas no texto da manchete

REGRAS PARA AS FRASES DE AÇÃO:
- Analise as marcas: {marcas_str}
- Para cada AÇÃO DISTINTA de uma marca no parágrafo, gere UMA frase
- A frase deve deixar claro: QUEM fez O QUÊ e para/com QUEM (quando relevante)
- PREFIRA nomes próprios a cargos ou descrições. Exemplo:
  ✅ "J&F fez repasse de R$ 25 milhões a Ibaneis Rocha"
  ❌ "J&F fez repasse ao escritório de advocacia do governador do Distrito Federal"
- As frases podem ser um pouco mais longas se necessário para serem precisas
- Não repita a mesma informação em frases diferentes
- Gere SOMENTE frases para marcas que realmente aparecem no parágrafo

PARÁGRAFO A ANALISAR:
{texto_para_prompt}

MARCAS A CONSIDERAR: {marcas_str}

Responda SOMENTE com o JSON. Nada mais."""

    try:
        resposta_bruta = _chamar_deepseek(
            prompt, chave_api, url_api,
            temperatura=TEMPERATURA_DEEPSEEK,
            max_tokens=MAX_TOKENS_PARAGRAFO
        )

        # Tentar extrair JSON da resposta
        # Às vezes o modelo inclui ```json ... ``` mesmo instruído a não incluir
        resposta_limpa = re.sub(r'```(?:json)?', '', resposta_bruta).strip().strip('`').strip()

        m = re.search(r'\{.*\}', resposta_limpa, flags=re.DOTALL)
        if not m:
            print(f"      ⚠️  Resposta da API não contém JSON. Resposta: {resposta_bruta[:200]}")
            return "Resumo não disponível", []

        obj = json.loads(m.group(0))
        manchete = str(obj.get('manchete', 'Sem manchete')).strip()
        acoes_raw = obj.get('acoes', [])
        acoes = [str(a).strip() for a in acoes_raw if str(a).strip()]

        return manchete, acoes

    except json.JSONDecodeError as e:
        print(f"      ⚠️  Erro ao parsear JSON da resposta DeepSeek: {e}")
        print(f"      Resposta recebida: {resposta_bruta[:300] if 'resposta_bruta' in dir() else 'N/A'}")
        return "Resumo não disponível", []
    except requests.exceptions.Timeout:
        print(f"      ⚠️  Timeout na chamada à API DeepSeek")
        return "Timeout na API", []
    except requests.exceptions.HTTPError as e:
        print(f"      ❌ Erro HTTP na API DeepSeek: {e}")
        return "Erro na API", []
    except Exception as e:
        print(f"      ❌ Erro inesperado ao chamar DeepSeek: {e}")
        return "Erro no processamento", []


# ==============================================================================
# GERAÇÃO DO DOCX CHK
# ==============================================================================

def _add_paragrafo_negrito(doc, texto, tamanho_pt=None, cor_rgb=None):
    """Adiciona parágrafo com texto em negrito ao documento."""
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    if tamanho_pt:
        run.font.size = Pt(tamanho_pt)
    if cor_rgb:
        run.font.color.rgb = RGBColor(*cor_rgb)
    return p


def _add_paragrafo_normal(doc, texto, tamanho_pt=None):
    """Adiciona parágrafo com texto normal ao documento."""
    p = doc.add_paragraph()
    run = p.add_run(texto)
    if tamanho_pt:
        run.font.size = Pt(tamanho_pt)
    return p


def gerar_docx_chk(arquivo_saida_chk, conteudo_processado):
    """
    Gera o arquivo DOCX CHK a partir da estrutura de conteúdo processado.

    conteudo_processado: lista de dicts com campos:
        tipo: 'cabecalho' | 'secao' | 'paragrafo_chk' | 'paragrafo_sem_marca'
        texto: str (texto original do parágrafo)
        dados: dict ou None
            Para 'paragrafo_chk': {'manchete': str, 'acoes': list, 'marcas': list}
    """
    doc = Document()

    # Configurar margens (opcional – mantém padrão se não configurado)
    from docx.shared import Inches
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    for item in conteudo_processado:
        tipo = item['tipo']
        texto = item['texto']
        dados = item.get('dados')

        if tipo == 'cabecalho':
            # Cabeçalho do relatório: negrito, tamanho maior
            _add_paragrafo_negrito(doc, texto, tamanho_pt=13)

        elif tipo == 'secao':
            # Título de seção: negrito, separador visual
            doc.add_paragraph("")  # Espaço antes
            _add_paragrafo_negrito(doc, texto, tamanho_pt=12, cor_rgb=(0, 70, 127))
            doc.add_paragraph("")  # Espaço depois

        elif tipo == 'paragrafo_chk':
            manchete = dados.get('manchete', 'Sem manchete')
            acoes = dados.get('acoes', [])

            # Manchete do parágrafo (bold, destaque)
            p_manchete = doc.add_paragraph()
            run_manchete = p_manchete.add_run(manchete)
            run_manchete.bold = True
            run_manchete.font.size = Pt(11)
            run_manchete.font.color.rgb = RGBColor(30, 30, 30)

            # Frases de ação (lista com bullet • )
            if acoes:
                for acao in acoes:
                    p_acao = doc.add_paragraph()
                    run_acao = p_acao.add_run(f"• {acao}")
                    run_acao.font.size = Pt(10)
            else:
                p_vazio = doc.add_paragraph()
                run_vazio = p_vazio.add_run("• (Nenhuma ação identificada pela IA)")
                run_vazio.italic = True
                run_vazio.font.size = Pt(10)

            # Separador entre blocos
            doc.add_paragraph("")

    doc.save(arquivo_saida_chk)


# ==============================================================================
# UPLOAD GOOGLE DRIVE (fallback autônomo, caso import de relatorio_ajustado_final falhe)
# ==============================================================================

def _upload_chk_drive(caminho_arquivo, pasta_id):
    """
    Upload do arquivo CHK para o Google Drive via Service Account.
    Espelho da função upload_para_google_drive do relatorio_ajustado_final.py.
    Usado como fallback caso o import direto não esteja disponível.
    """
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload

    SERVICE_ACCOUNT_FILE = 'service_account.json'
    SCOPES = ['https://www.googleapis.com/auth/drive']

    creds = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES
    )
    service = build('drive', 'v3', credentials=creds)

    file_metadata = {
        'name': os.path.basename(caminho_arquivo),
        'parents': [pasta_id]
    }
    media = MediaFileUpload(
        caminho_arquivo,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    uploaded = service.files().create(
        body=file_metadata,
        media_body=media,
        fields='id',
        supportsAllDrives=True
    ).execute()
    print(f"  Upload CHK concluído com ID: {uploaded.get('id')}")


# ==============================================================================
# FUNÇÃO PRINCIPAL DE PROCESSAMENTO
# ==============================================================================

def processar_verificacao_chk(arquivo_docx, pasta_id_drive=None, pasta_destino='/app/output'):
    """
    Ponto de entrada principal.

    Lê o DOCX gerado, verifica a presença da seção de marcas,
    processa cada parágrafo relevante com DeepSeek e gera o arquivo _CHK.

    Args:
        arquivo_docx (str): Caminho do arquivo DOCX gerado pelo relatorio_ajustado_final.py
        pasta_id_drive (str|None): ID da pasta no Google Drive para upload do CHK.
                                   Deve ser o mesmo usado para o arquivo principal.
        pasta_destino (str): Caminho local de destino (volume Docker mapeado para a rede).

    Returns:
        bool: True se processamento concluído com sucesso, False caso contrário.
    """
    separador = "=" * 65
    print(f"\n{separador}")
    print(f"  🔍  VERIFICAÇÃO CHK — ANÁLISE DE MARCAS")
    print(separador)
    print(f"  Arquivo de entrada : {arquivo_docx}")

    # ------------------------------------------------------------------
    # 1. Verificar existência do arquivo
    # ------------------------------------------------------------------
    if not os.path.exists(arquivo_docx):
        print(f"\n  ❌ Arquivo não encontrado: {arquivo_docx}")
        print(f"{separador}\n")
        return False

    # ------------------------------------------------------------------
    # 2. Abrir documento e verificar condição de entrada
    # ------------------------------------------------------------------
    print(f"\n  📋 [1/6] Abrindo e verificando estrutura do documento...")
    try:
        doc = Document(arquivo_docx)
    except Exception as e:
        print(f"  ❌ Erro ao abrir o documento DOCX: {e}")
        return False

    textos_paragrafos = [p.text.strip() for p in doc.paragraphs]
    linhas_nao_vazias = [t for t in textos_paragrafos if t]

    # Verificar nas primeiras N linhas não-vazias
    primeiras = linhas_nao_vazias[:LINHAS_VERIFICACAO]
    contem_secao_marcas = any('--- NOTÍCIAS DE MARCAS ---' in linha for linha in primeiras)

    if not contem_secao_marcas:
        print(f"\n  ⚠️  '--- NOTÍCIAS DE MARCAS ---' não encontrada nas primeiras "
              f"{LINHAS_VERIFICACAO} linhas não-vazias do documento.")
        print(f"  ⏭️  Processamento CHK ignorado para este arquivo.")
        print(f"{separador}\n")
        return False

    print(f"  ✅ Seção '--- NOTÍCIAS DE MARCAS ---' encontrada. Prosseguindo.")

    # ------------------------------------------------------------------
    # 3. Obter credenciais DeepSeek
    # ------------------------------------------------------------------
    print(f"\n  🔑 [2/6] Obtendo credenciais DeepSeek...")
    try:
        chave_api = _obter_chave_deepseek()
        url_api = _obter_url_deepseek()
        print(f"  ✅ Chave API obtida. URL: {url_api}")
    except Exception as e:
        print(f"  ❌ Falha ao obter credenciais DeepSeek: {e}")
        print(f"{separador}\n")
        return False

    # ------------------------------------------------------------------
    # 4. Percorrer parágrafos e classificar conteúdo
    # ------------------------------------------------------------------
    print(f"\n  📄 [3/6] Varrendo parágrafos do documento...")

    conteudo_processado = []
    secao_atual = None          # Seção em que estamos no momento
    antes_da_primeira_secao = True  # Flag para capturar cabeçalho do relatório

    total_paragrafos_com_marca = 0
    total_paragrafos_sem_marca = 0
    total_marcas_detectadas = 0

    for paragrafo in doc.paragraphs:
        texto = paragrafo.text.strip()
        if not texto:
            continue

        # Verificar se o parágrafo é um título de seção (padrão "--- ... ---")
        eh_titulo_secao = bool(re.match(r'^---\s*.+\s*---\s*$', texto))
        eh_secao_relevante = any(sec in texto for sec in SECOES_RELEVANTES)

        if eh_titulo_secao:
            if eh_secao_relevante:
                antes_da_primeira_secao = False
                for sec in SECOES_RELEVANTES:
                    if sec in texto:
                        secao_atual = sec
                        break
                conteudo_processado.append({
                    'tipo': 'secao',
                    'texto': texto,
                    'dados': None
                })
                print(f"  📌 Seção relevante encontrada: {secao_atual}")
            else:
                # Seção não monitorada — desativa o processamento de parágrafos
                print(f"  ⏭️  Seção ignorada: {texto}")
                secao_atual = None
            continue

        # Capturar linhas de cabeçalho (antes da primeira seção relevante)
        if antes_da_primeira_secao:
            conteudo_processado.append({
                'tipo': 'cabecalho',
                'texto': texto,
                'dados': None
            })
            continue

        # Se estamos dentro de uma seção relevante, processar
        if secao_atual in SECOES_RELEVANTES:
            marcas = detectar_marcas_no_paragrafo(texto)

            if marcas:
                total_paragrafos_com_marca += 1
                total_marcas_detectadas += len(marcas)
                conteudo_processado.append({
                    'tipo': 'paragrafo_chk',
                    'texto': texto,
                    'dados': {
                        'manchete': None,       # Será preenchido no passo 4
                        'acoes': [],            # Será preenchido no passo 4
                        'marcas': marcas
                    }
                })
                print(f"  ✔  Parágrafo {total_paragrafos_com_marca}: "
                      f"{len(marcas)} marca(s) → {', '.join(marcas)}")
            else:
                total_paragrafos_sem_marca += 1
                # Parágrafos sem marcas monitoradas são ignorados no relatório CHK

    print(f"\n  📊 Varredura concluída:")
    print(f"     Parágrafos COM marcas  : {total_paragrafos_com_marca}")
    print(f"     Parágrafos sem marcas  : {total_paragrafos_sem_marca}")
    print(f"     Total de marcas detectadas : {total_marcas_detectadas}")

    if total_paragrafos_com_marca == 0:
        print(f"\n  ℹ️  Nenhum parágrafo com marcas monitoradas encontrado.")
        print(f"  ⏭️  Arquivo CHK não será gerado (sem conteúdo a processar).")
        print(f"{separador}\n")
        return False

    # ------------------------------------------------------------------
    # 5. Chamar DeepSeek para cada parágrafo com marcas
    # ------------------------------------------------------------------
    print(f"\n  🤖 [4/6] Processando {total_paragrafos_com_marca} parágrafo(s) com DeepSeek...")

    paragrafos_chk = [item for item in conteudo_processado if item['tipo'] == 'paragrafo_chk']
    erros_api = 0

    for idx, item in enumerate(paragrafos_chk, start=1):
        texto_paragrafo = item['texto']
        marcas = item['dados']['marcas']

        print(f"\n  [{idx}/{total_paragrafos_com_marca}] Marcas: {', '.join(marcas)}")
        print(f"     Texto: {texto_paragrafo[:100]}{'...' if len(texto_paragrafo) > 100 else ''}")

        manchete, acoes = gerar_manchete_e_acoes(
            texto_paragrafo, marcas, chave_api, url_api
        )

        item['dados']['manchete'] = manchete
        item['dados']['acoes'] = acoes

        print(f"     ✅ Manchete  : {manchete}")
        print(f"     ✅ Ações ({len(acoes)}): {'; '.join(acoes[:2])}{'...' if len(acoes) > 2 else ''}")

        if manchete in ("Resumo não disponível", "Timeout na API", "Erro na API", "Erro no processamento"):
            erros_api += 1

        # Pausa entre chamadas para evitar rate limit
        if idx < total_paragrafos_com_marca:
            time.sleep(PAUSA_ENTRE_CHAMADAS)

    print(f"\n  📊 Chamadas DeepSeek concluídas. Erros: {erros_api}/{total_paragrafos_com_marca}")

    # ------------------------------------------------------------------
    # 6. Definir nome e caminho do arquivo CHK
    # ------------------------------------------------------------------
    print(f"\n  📁 [5/6] Definindo nome do arquivo CHK...")

    # arquivo_docx pode ser "output/Destaques do dia - J&F_20260323_0354.docx"
    # CHK será  "output/Destaques do dia - J&F_20260323_0354_CHK.docx"
    arquivo_chk = arquivo_docx.replace('.docx', '_CHK.docx')

    print(f"  ✅ Arquivo CHK: {arquivo_chk}")

    # ------------------------------------------------------------------
    # 7. Gerar o documento DOCX CHK
    # ------------------------------------------------------------------
    print(f"\n  💾 [6/6] Gerando arquivo DOCX CHK...")

    try:
        gerar_docx_chk(arquivo_chk, conteudo_processado)
    except Exception as e:
        print(f"  ❌ Erro ao gerar o documento CHK: {e}")
        import traceback
        traceback.print_exc()
        print(f"{separador}\n")
        return False

    # Verificar se o arquivo foi realmente criado
    if os.path.exists(arquivo_chk):
        tamanho_kb = os.path.getsize(arquivo_chk) / 1024
        print(f"  ✅ Arquivo CHK gerado com sucesso!")
        print(f"     Caminho : {arquivo_chk}")
        print(f"     Tamanho : {tamanho_kb:.1f} KB")
    else:
        print(f"  ❌ Arquivo CHK não encontrado após tentativa de gravação.")
        print(f"{separador}\n")
        return False

    # ------------------------------------------------------------------
    # 8. Upload do CHK para o Google Drive (mesma pasta do arquivo principal)
    # ------------------------------------------------------------------
    if pasta_id_drive:
        print(f"\n  ☁️  Fazendo upload do CHK para o Google Drive...")
        try:
            from relatorio_ajustado_final import upload_para_google_drive
            upload_para_google_drive(arquivo_chk, os.path.basename(arquivo_chk), pasta_id_drive)
            print(f"  ✅ Upload do CHK concluído.")
        except ImportError:
            # Fallback: reimplementar o upload diretamente aqui
            print(f"  ⚠️  Não foi possível importar upload_para_google_drive. "
                  f"Tentando upload direto...")
            try:
                _upload_chk_drive(arquivo_chk, pasta_id_drive)
                print(f"  ✅ Upload do CHK concluído (via fallback).")
            except Exception as e:
                print(f"  ⚠️  Upload do CHK falhou: {e}. "
                      f"O arquivo foi salvo localmente mas pode não aparecer no Drive.")
        except Exception as e:
            print(f"  ⚠️  Erro no upload do CHK: {e}. "
                  f"O arquivo foi salvo localmente mas pode não aparecer no Drive.")
    else:
        print(f"  ℹ️  pasta_id_drive não informado — upload do CHK ignorado.")

    # ------------------------------------------------------------------
    # 9. Cópia para pasta de rede local (volume Docker)
    # ------------------------------------------------------------------
    if os.path.isdir(pasta_destino):
        import shutil
        destino_chk = os.path.join(pasta_destino, os.path.basename(arquivo_chk))
        try:
            if not (os.path.exists(destino_chk) and os.path.samefile(arquivo_chk, destino_chk)):
                shutil.copy2(arquivo_chk, destino_chk)
                print(f"  ✅ CHK também copiado para pasta de rede: {destino_chk}")
            else:
                print(f"  ℹ️  CHK já está na pasta de destino (mesmo arquivo).")
        except Exception as e:
            print(f"  ⚠️  Erro ao copiar CHK para pasta de rede: {e}")

    print(f"\n{separador}")
    print(f"  ✅  VERIFICAÇÃO CHK CONCLUÍDA")
    print(f"  Parágrafos processados : {total_paragrafos_com_marca}")
    print(f"  Erros na API           : {erros_api}")
    print(f"  Arquivo gerado         : {os.path.basename(arquivo_chk)}")
    print(f"{separador}\n")

    return True


# ==============================================================================
# EXECUÇÃO STANDALONE (para testes)
# ==============================================================================

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Uso: python verificador_marcas_chk.py <caminho_do_arquivo.docx>")
        print("Exemplo: python verificador_marcas_chk.py output/Destaques\\ do\\ dia\\ -\\ J&F_20260323_0354.docx")
        sys.exit(1)

    arquivo = sys.argv[1]
    sucesso = processar_verificacao_chk(arquivo)
    sys.exit(0 if sucesso else 1)