# Versão Preliminar do Relatório

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import re
import time
from config import arq_resumo_final, marca1, marca2
import requests
import os
from datetime import datetime

# IMPORTAR O SISTEMA DE ENCURTAMENTO OTIMIZADO
from encurtador_urls import GerenciadorURLs, testar_servicos_encurtamento


def processar_editoriais_integrado(final_df_editorial, document, gerenciador_urls):
    """
    Processa editoriais de forma integrada com gerenciador de URLs
    """
    print(f"\nProcessando {len(final_df_editorial)} editoriais...")
    
    if final_df_editorial.empty:
        print("Nenhum editorial encontrado.")
        return
    
    # Adicionar cabeçalho da seção
    document.add_paragraph("")
    document.add_paragraph("*EDITORIAIS*")
    document.add_paragraph("")

    # Debug: verificar estrutura do DataFrame
    print(f"Colunas disponíveis: {list(final_df_editorial.columns)}")

    for index, row_editorial in final_df_editorial.iterrows():
        try:
            # Tentar diferentes nomes de colunas possíveis
            w_veiculo_editorial = (
                row_editorial.get('Veiculo') or 
                row_editorial.get('Veículo') or 
                row_editorial.get('veiculo') or 
                'Veículo Desconhecido'
            )
            
            w_titulo_editorial = (
                row_editorial.get('Titulo') or 
                row_editorial.get('Título') or 
                row_editorial.get('titulo') or 
                row_editorial.get('Conteudo') or 
                row_editorial.get('Conteúdo') or 
                'Título não disponível'
            )
            
            w_url_editorial = (
                row_editorial.get('UrlVisualizacao') or 
                row_editorial.get('Url') or 
                row_editorial.get('URL') or 
                row_editorial.get('Link') or 
                'URL Não Disponível'
            )
            
            print(f"\nEditorial {index + 1}:")
            print(f"  Veículo: {w_veiculo_editorial}")
            print(f"  Título: {w_titulo_editorial[:50]}...")
            print(f"  URL: {w_url_editorial}")
            
            # Adicionar informações básicas
            document.add_paragraph(f"{w_veiculo_editorial}: {w_titulo_editorial}")
            
            # Encurtamento seguro da URL usando o gerenciador
            if w_url_editorial and w_url_editorial != 'URL Não Disponível':
                short_url_editorial = gerenciador_urls.obter_url_curta(w_url_editorial)
            else:
                short_url_editorial = w_url_editorial
                print(f"URL inválida para editorial {index + 1}")
            
            document.add_paragraph(short_url_editorial)
            document.add_paragraph("*")
            
            print(f"Editorial {index + 1} processado ✅")
            
        except Exception as e:
            print(f"Erro ao processar editorial {index + 1}: {str(e)}")
            document.add_paragraph(f"Editorial {index + 1}: Erro no processamento")
            document.add_paragraph("URL não disponível")
            document.add_paragraph("*")
            continue


def gerar_versao_preliminar(final_df_small_marca, final_df_small_marca_irrelevantes, df_resumo_marca, df_resumo_marca_irrelevantes, final_df_marca, df_resumo_setor, final_df_setor, final_df_editorial, final_df_SPECIALS_small):
    
    # ========================================================================
    # INICIALIZAR GERENCIADOR DE URLs COM CACHE
    # ========================================================================
    print("\n" + "="*80)
    print("🔗 INICIALIZANDO SISTEMA DE ENCURTAMENTO DE URLs")
    print("="*80)
    
    gerenciador_urls = GerenciadorURLs('dados/cache_urls.json')
    
    print("="*80 + "\n")
    
    # ========================================================================
    # INICIALIZAR DOCUMENTO DOCX
    # ========================================================================
    
    document = Document()

    # Configurar o estilo Normal
    styles = document.styles
    style = styles['Normal']
    font = style.font
    font.name = 'Calibri'
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Adicionar título e linha em branco
    document.add_paragraph("DESTAQUES DO DIA J&F")
    document.add_paragraph("")
    document.add_paragraph("--- NOTÍCIAS DE MARCAS ---")
    document.add_paragraph("")

    # ========================================================================
    # SEÇÃO: RESUMOS DE MARCA (RELEVANTES)
    # ========================================================================
    
    print(f"Processando {len(df_resumo_marca)} resumos de Marca...")
    
    for index, row_marca in df_resumo_marca.iterrows():
        group_string = ""

        if 'Ids' not in row_marca or pd.isna(row_marca['Ids']):
            print(f"Aviso: Linha {index} sem IDs. Pulando.")
            continue

        for news_id_str in str(row_marca['Ids']).split(','):
            try:
                news_id = int(news_id_str.strip())
            except ValueError:
                print(f"Aviso: ID inválido '{news_id_str}'. Pulando.")
                continue

            news_info_marca = final_df_marca[final_df_marca['Id'] == news_id]
            if news_info_marca.empty:
                print(f"Aviso: ID {news_id} não encontrado. Pulando.")
                continue

            news_info_marca = news_info_marca.iloc[0]

            w_veiculo_marca = news_info_marca['Veiculo']
            w_url_marca = news_info_marca['UrlVisualizacao']

            # ENCURTAR URL USANDO GERENCIADOR (COM CACHE)
            short_url_marca = gerenciador_urls.obter_url_curta(w_url_marca)

            # Atualizar DataFrame com URL curta
            if news_id in final_df_small_marca['Id'].values:
                final_df_small_marca.loc[final_df_small_marca['Id'] == news_id, 'ShortURL'] = short_url_marca

            # Identificar tipo "Special"
            special_type = ""
            special_info = final_df_SPECIALS_small[final_df_SPECIALS_small['Id'] == news_id]

            if not special_info.empty:
                canais_special = special_info.iloc[0]['Canais']
                if isinstance(canais_special, list):
                    canais_str = ', '.join(map(str, canais_special))
                else:
                    canais_str = str(canais_special)

                if "Editoriais" in canais_str:
                    special_type = "Editorial"
                elif "Colunistas" in canais_str:
                    special_type = "Colunista"
                elif "1ª Página" in canais_str:
                    special_type = "Capa"

            # Formatar string
            if special_type:
                group_string += f"{w_veiculo_marca} ({special_type} - {short_url_marca}), "
            else:
                group_string += f"{w_veiculo_marca} ({short_url_marca}), "

        # Limpar e adicionar resumo
        group_string = group_string.rstrip(', ')
        group_string += "\n"

        if 'Resumo' not in row_marca or pd.isna(row_marca['Resumo']):
            resumo_limpo = "[Resumo não disponível]"
        else:
            resumo_limpo = str(row_marca['Resumo'])

        group_string += resumo_limpo

        document.add_paragraph(group_string)
        document.add_paragraph("")

    # ========================================================================
    # SEÇÃO: CITAÇÕES (NOTÍCIAS IRRELEVANTES)
    # ========================================================================

    document.add_paragraph("")
    document.add_paragraph("--- CITAÇÕES ---")
    document.add_paragraph("")
    
    print(f"Processando {len(df_resumo_marca_irrelevantes)} citações...")
    
    for index, row_marca in df_resumo_marca_irrelevantes.iterrows():
        group_string = ""

        if 'Ids' not in row_marca or pd.isna(row_marca['Ids']):
            print(f"Aviso: Linha {index} sem IDs. Pulando.")
            continue

        for news_id_str in str(row_marca['Ids']).split(','):
            try:
                news_id = int(news_id_str.strip())
            except ValueError:
                continue

            news_info_marca = final_df_marca[final_df_marca['Id'] == news_id]
            if news_info_marca.empty:
                continue

            news_info_marca = news_info_marca.iloc[0]

            w_veiculo_marca = news_info_marca['Veiculo']
            w_url_marca = news_info_marca['UrlVisualizacao']

            # ENCURTAR URL USANDO GERENCIADOR (COM CACHE)
            short_url_marca = gerenciador_urls.obter_url_curta(w_url_marca)

            # Atualizar DataFrame
            if news_id in final_df_small_marca_irrelevantes['Id'].values:
                final_df_small_marca_irrelevantes.loc[final_df_small_marca_irrelevantes['Id'] == news_id, 'ShortURL'] = short_url_marca

            # Identificar tipo "Special"
            special_type = ""
            special_info = final_df_SPECIALS_small[final_df_SPECIALS_small['Id'] == news_id]

            if not special_info.empty:
                canais_special = special_info.iloc[0]['Canais']
                if isinstance(canais_special, list):
                    canais_str = ', '.join(map(str, canais_special))
                else:
                    canais_str = str(canais_special)

                if "Editoriais" in canais_str:
                    special_type = "Editorial"
                elif "Colunistas" in canais_str:
                    special_type = "Colunista"
                elif "1ª Página" in canais_str:
                    special_type = "Capa"

            if special_type:
                group_string += f"{w_veiculo_marca} ({special_type} - {short_url_marca}), "
            else:
                group_string += f"{w_veiculo_marca} ({short_url_marca}), "

        group_string = group_string.rstrip(', ')
        group_string += "\n"

        if 'Resumo' not in row_marca or pd.isna(row_marca['Resumo']):
            resumo_limpo = "[Resumo não disponível]"
        else:
            resumo_limpo = str(row_marca['Resumo'])

        group_string += resumo_limpo

        document.add_paragraph(group_string)
        document.add_paragraph("")

    # ========================================================================
    # SEÇÃO: LINKS DAS NOTÍCIAS POR MARCAS
    # ========================================================================

    final_df_small_marca_combined = pd.concat([final_df_small_marca, final_df_small_marca_irrelevantes], ignore_index=True)

    print(f"\nVerificando links por marcas: {len(final_df_small_marca_combined)} notícias")

    if not final_df_small_marca_combined.empty and 'Canais' in final_df_small_marca_combined.columns:
        try:
            order = [marca1, marca2] + [marca for marca in final_df_small_marca_combined['Canais'].unique() if marca not in (marca1, marca2)]
            final_df_small_marca_sorted = final_df_small_marca_combined.sort_values(by=['Canais'], key=lambda x: pd.Categorical(x, categories=order, ordered=True))
            print("Ordenação personalizada aplicada.")
        except Exception as e:
            print(f"Aviso: Falha na ordenação ({e}). Usando ordem padrão.")
            final_df_small_marca_sorted = final_df_small_marca_combined.sort_values(by=['Canais'])

        document.add_paragraph("")
        document.add_paragraph("--- LINKS DAS NOTÍCIAS DE MARCA ---")
        document.add_paragraph("")

        current_marca = None

        for index, row_small_marca in final_df_small_marca_sorted.iterrows():
            marca = row_small_marca['Canais']
            if marca != current_marca:
                if current_marca is not None:
                    document.add_paragraph("")
                document.add_paragraph(f"*{marca}*")
                current_marca = marca

            news_id_small = row_small_marca['Id']
            original_row_marca = final_df_marca[final_df_marca['Id'] == news_id_small]
            if original_row_marca.empty:
                continue

            original_row_marca = original_row_marca.iloc[0]

            veiculo = original_row_marca['Veiculo']
            titulo = original_row_marca['Titulo']
            canais_commodities = original_row_marca.get('CanaisCommodities', '')

            document.add_paragraph(f"{veiculo}: {titulo}")

            short_url = row_small_marca.get('ShortURL')
            if pd.isna(short_url) or not short_url:
                short_url = original_row_marca.get('UrlVisualizacao', 'URL Não Encontrada')

            prefix = "Coluna - " if "Colunistas" in str(canais_commodities) else ""
            document.add_paragraph(f"{prefix}{short_url}")

            document.add_paragraph("*")

    # ========================================================================
    # SEÇÃO: RESUMOS DE SETOR
    # ========================================================================

    print(f"\nVerificando resumos de setor: {len(df_resumo_setor)} notícias")
    
    if not df_resumo_setor.empty:
        document.add_paragraph("")
        document.add_paragraph("--- NOTÍCIAS DE SETOR ---")
        document.add_paragraph("")

    current_tema = None

    for index, row_setor in df_resumo_setor.iterrows():
        if 'Tema' not in row_setor or pd.isna(row_setor['Tema']):
            continue

        tema = row_setor['Tema']

        if tema != current_tema:
            if current_tema is not None:
                document.add_paragraph("")
            document.add_paragraph(f"*{tema}*")
            current_tema = tema

        if 'Id' not in row_setor or pd.isna(row_setor['Id']):
            continue

        try:
            news_id = int(str(row_setor['Id']).strip())
        except ValueError:
            continue

        news_info_setor = final_df_setor[final_df_setor['Id'] == news_id]
        if news_info_setor.empty:
            continue
            
        news_info_setor = news_info_setor.iloc[0]

        w_veiculo_setor = news_info_setor['Veiculo']
        w_titulo_setor = news_info_setor['Titulo']
        w_url_setor = news_info_setor['UrlVisualizacao']

        document.add_paragraph(f"{w_veiculo_setor}: {w_titulo_setor}")

        if 'Resumo' not in row_setor or pd.isna(row_setor['Resumo']):
            document.add_paragraph("[Resumo não disponível]")
        else:
            resumo_bruto = str(row_setor['Resumo'])
            document.add_paragraph(resumo_bruto)

        # ENCURTAR URL USANDO GERENCIADOR (COM CACHE)
        short_url_setor = gerenciador_urls.obter_url_curta(w_url_setor)

        document.add_paragraph(short_url_setor)
        document.add_paragraph("*")

    # ========================================================================
    # SEÇÃO: EDITORIAIS
    # ========================================================================

    processar_editoriais_integrado(final_df_editorial, document, gerenciador_urls)

    # ========================================================================
    # PROCESSAMENTO DO DOCUMENTO ANTES DA GRAVAÇÃO
    # ========================================================================

    print("\nProcessando documento para remover padrões indesejados...")

    paragraphs_to_remove_indices = []

    pattern_line_start_paren_end_palavras = r"^\s*\*\s*\(.*?palavras\)\s*$"
    pattern_line_start_resumo = r"^\s*\*Resumo.*$"
    pattern_parenthesized_palavras_general = r"\s*[\*\s]*\(.*?\s*palavras\s*\)[\s\:\*]*"
    pattern_specific_resumo_prefixes = r"^\s*[\*\s]*Resumo\s*[:\*\s]*"

    for i, paragraph in enumerate(document.paragraphs):
        texto = paragraph.text

        if re.fullmatch(pattern_line_start_paren_end_palavras, texto, re.IGNORECASE):
            paragraphs_to_remove_indices.append(i)
            continue

        if re.fullmatch(pattern_line_start_resumo, texto, re.IGNORECASE):
            paragraphs_to_remove_indices.append(i)
            continue

        new_text = re.sub(pattern_parenthesized_palavras_general, '', texto, flags=re.IGNORECASE).strip()
        new_text = re.sub(pattern_specific_resumo_prefixes, '', new_text, flags=re.IGNORECASE).strip()

        if new_text != texto.strip():
            paragraph.text = new_text

        if "**" in paragraph.text:
            paragraph.text = paragraph.text.replace("**", "*")

    # Remover parágrafos marcados
    for i in sorted(paragraphs_to_remove_indices, reverse=True):
        p = document.paragraphs[i]._element
        p.getparent().remove(p)

    print(f"Total de parágrafos removidos: {len(paragraphs_to_remove_indices)}")

    # ========================================================================
    # EXIBIR ESTATÍSTICAS DO CACHE
    # ========================================================================
    
    gerenciador_urls.estatisticas()

    # ========================================================================
    # SALVAR DOCUMENTO
    # ========================================================================

    document.save(arq_resumo_final)
    print(f"\n✅ Arquivo DOCX salvo: {arq_resumo_final}")