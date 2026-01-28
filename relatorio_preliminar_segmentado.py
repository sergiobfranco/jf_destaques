# Versão Segmentada do Relatório - v2
# Atualização: Resumos de Marca e Citações agora ficam na mesma linha das publicações

import pandas as pd
import pyshorteners
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import re
import time
from config import arq_resumo_final, marca1, marca2
import requests
import os
import traceback
from datetime import datetime

# 1. FUNÇÃO AUXILIAR PARA ENCURTAMENTO DE URL
def encurtar_url_seguro(url_original, max_tentativas=3, delay=2):
    """
    Função auxiliar para encurtar URLs de forma segura com tratamento de erro robusto
    """
    import pyshorteners
    import requests
    import time
    
    if not url_original or pd.isna(url_original) or str(url_original).strip() == '':
        print(f"URL vazia ou inválida: {url_original}")
        return str(url_original) if url_original else "URL não disponível"
    
    url_str = str(url_original).strip()
    
    if not url_str.startswith(('http://', 'https://')):
        print(f"URL não tem protocolo válido: {url_str}")
        return url_str
    
    s = pyshorteners.Shortener()
    
    for tentativa in range(max_tentativas):
        try:
            short_url = s.tinyurl.short(url_str)
            return short_url
            
        except Exception as e:
            erro_str = str(e)
            print(f"Tentativa {tentativa + 1}/{max_tentativas} - Erro TinyURL: {erro_str}")
            
            if tentativa == max_tentativas - 1:
                try:
                    short_url = s.isgd.short(url_str)
                    print(f"URL encurtada com serviço alternativo (is.gd)")
                    return short_url
                except Exception as e2:
                    print(f"Erro também no serviço alternativo: {str(e2)}")
                    break
            
            if tentativa < max_tentativas - 1:
                time.sleep(delay)
    
    print(f"Falha em todos os serviços de encurtamento. Usando URL original.")
    return url_str

def processar_editoriais_integrado(final_df_editorial, document, opcao_selecionada=1):
    """
    Processa editoriais de forma integrada (substitui toda a seção de editoriais)
    """
    print(f"\nProcessando {len(final_df_editorial)} editoriais...")
    
    if final_df_editorial.empty:
        print("Nenhum editorial encontrado.")
        return

    print(f"Colunas disponíveis em final_df_editorial: {list(final_df_editorial.columns)}")
    print(f"Primeiras linhas do DataFrame:")
    print(final_df_editorial.head())

    for index, row_editorial in final_df_editorial.iterrows():
        try:
            w_veiculo_editorial = (
                row_editorial.get('Veiculo') or 
                row_editorial.get('Veículo') or 
                row_editorial.get('veiculo') or 
                'Veículo Desconhecido'
            )
            w_veiculo_editorial = w_veiculo_editorial.title()
            
            w_titulo_editorial = (
                row_editorial.get('Titulo') or 
                row_editorial.get('Título') or 
                row_editorial.get('titulo') or 
                row_editorial.get('Conteudo') or 
                row_editorial.get('Conteúdo') or 
                'Título não disponível'
            )
            
            w_url_editorial = (
                row_editorial.get('UrlVisualizacao') or 
                row_editorial.get('Url') or 
                row_editorial.get('URL') or 
                row_editorial.get('Link') or 
                'URL Não Disponível'
            )
            
            print(f"\nEditorial {index + 1}:")
            print(f"  Veículo: {w_veiculo_editorial}")
            print(f"  Título: {w_titulo_editorial}")
            print(f"  URL: {w_url_editorial}")
            
            document.add_paragraph(f"{w_veiculo_editorial}: {w_titulo_editorial}")
            
            if w_url_editorial and w_url_editorial != 'URL Não Disponível':
                short_url_editorial = encurtar_url_seguro(w_url_editorial, max_tentativas=3, delay=1)
            else:
                short_url_editorial = w_url_editorial
                print(f"URL inválida para editorial {index + 1}, usando valor original")
            
            document.add_paragraph(short_url_editorial)
            document.add_paragraph("*")
            
            print(f"Editorial {index + 1} processado com sucesso")
            
        except Exception as e:
            print(f"Erro ao processar editorial {index + 1}: {str(e)}")
            print(f"Dados do editorial: {dict(row_editorial)}")
            document.add_paragraph(f"Editorial {index + 1}: Erro no processamento")
            document.add_paragraph("URL não disponível")
            document.add_paragraph("*")
            continue

def debug_dataframe_structure(df, nome_df):
    """
    Função auxiliar para debugar a estrutura dos DataFrames
    """
    print(f"\n=== DEBUG: Estrutura do {nome_df} ===")
    print(f"Shape: {df.shape}")
    print(f"Colunas: {list(df.columns)}")
    if not df.empty:
        print(f"Primeiras 2 linhas:")
        for i in range(min(2, len(df))):
            print(f"Linha {i}: {dict(df.iloc[i])}")
    else:
        print("DataFrame vazio")
    print("=" * 50)

def obter_nome_veiculo(codigo_veiculo):
    """
    Retorna o nome do veículo baseado no código
    """
    nomes_veiculos = {
        675: "O ESTADO DE S.PAULO/SÃO PAULO",
        10459: "VALOR ECONÔMICO / SÃO PAULO", 
        331: "FOLHA DE S.PAULO / SÃO PAULO",
        682: "O GLOBO / RIO DE JANEIRO"
    }
    nome = nomes_veiculos.get(codigo_veiculo, f"VEÍCULO CÓDIGO {codigo_veiculo}")
    return nome.title()

def gerar_versao_preliminar(final_df_small_marca, final_df_small_marca_irrelevantes, df_resumo_marca, 
                          df_resumo_marca_irrelevantes, final_df_marca, df_resumo_setor, final_df_setor, 
                          final_df_editorial, final_df_SPECIALS_small, opcao_selecionada=1, codigo_veiculo=None):
    """
    Gera versão preliminar do relatório baseada na opção selecionada
    
    Args:
        opcao_selecionada (int): Opção selecionada (1-7)
        codigo_veiculo (int): Código do veículo (para opções 4-7)
    """
    
    # Determinar quais seções executar baseado na opção selecionada
    executar_marcas = opcao_selecionada in [1, 2]
    executar_setor = opcao_selecionada in [1, 3, 4, 5, 6, 7]
    executar_editoriais = opcao_selecionada in [1, 2, 3]
    
    print(f"\n=== CONFIGURAÇÃO DE PROCESSAMENTO ===")
    print(f"Opção selecionada: {opcao_selecionada}")
    print(f"Código do veículo: {codigo_veiculo}")
    print(f"Executar MARCAS: {executar_marcas}")
    print(f"Executar SETOR: {executar_setor}")
    print(f"Executar EDITORIAIS: {executar_editoriais}")
    print("=" * 50)
    
    # 0. SEÇÃO DE CONFIGURAÇÕES INICIAIS (SEMPRE EXECUTADA)
    print("\n=== SEÇÃO 0: CONFIGURAÇÕES INICIAIS ===")
    
    debug_dataframe_structure(final_df_editorial, "final_df_editorial")

    document = Document()

    styles = document.styles
    style = styles['Normal']
    font = style.font
    font.name = 'Calibri'
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    document.add_paragraph("DESTAQUES DO DIA J&F")
    document.add_paragraph("")
    
    # 1. SEÇÃO DE MARCAS (CONDICIONAL)
    if executar_marcas:
        print("\n=== SEÇÃO 1: MARCAS ===")
        
        document.add_paragraph("--- NOTÍCIAS DE MARCAS ---")
        document.add_paragraph("")

        print(f"Processando {len(df_resumo_marca)} resumos de Marca...")
        for index, row_marca in df_resumo_marca.iterrows():
            group_string = ""

            if 'Ids' not in row_marca or pd.isna(row_marca['Ids']):
                print(f"Aviso: Linha {index} no df_resumo_marca não tem IDs válidos. Pulando.")
                continue

            for news_id_str in str(row_marca['Ids']).split(','):
                try:
                    news_id = int(news_id_str.strip())
                except ValueError:
                    print(f"Aviso: Não foi possível converter ID '{news_id_str}' para inteiro na linha {index} de df_resumo_marca. Pulando.")
                    continue

                news_info_marca = final_df_marca[final_df_marca['Id'] == news_id]
                if news_info_marca.empty:
                    print(f"Aviso: ID {news_id} não encontrado em final_df_marca para resumo de Marca. Pulando.")
                    continue

                news_info_marca = news_info_marca.iloc[0]
                w_veiculo_marca = news_info_marca['Veiculo'].title()
                w_url_marca = news_info_marca['UrlVisualizacao']

                short_url_marca = encurtar_url_seguro(w_url_marca, max_tentativas=3, delay=1)

                if news_id in final_df_small_marca['Id'].values:
                    final_df_small_marca.loc[final_df_small_marca['Id'] == news_id, 'ShortURL'] = short_url_marca
                else:
                    print(f"Aviso: ID {news_id} não encontrado em final_df_small_marca para adicionar ShortURL.")

                special_type = ""
                special_info = final_df_SPECIALS_small[final_df_SPECIALS_small['Id'] == news_id]

                if not special_info.empty:
                    canais_special = special_info.iloc[0]['Canais']
                    if isinstance(canais_special, list):
                        canais_str = ', '.join(map(str, canais_special))
                    else:
                        canais_str = str(canais_special)

                    if "Editoriais" in canais_str:
                        special_type = "Editorial"
                    elif "Colunistas" in canais_str:
                        special_type = "Colunista"
                    elif "1ª Página" in canais_str:
                        special_type = "Capa"

                if special_type:
                    group_string += f"{w_veiculo_marca} ({special_type} - {short_url_marca}), "
                else:
                    group_string += f"{w_veiculo_marca} ({short_url_marca}), "

            # ========== CORREÇÃO APLICADA AQUI ==========
            # Limpar vírgula final e adicionar espaço (SEM \n)
            group_string = group_string.rstrip(', ')

            # Adicionar resumo na mesma linha
            if 'Resumo' not in row_marca or pd.isna(row_marca['Resumo']):
                print(f"Aviso: Linha {index} no df_resumo_marca não tem Resumo. Adicionando placeholder.")
                resumo_limpo = "[Resumo não disponível]"
            else:
                resumo_limpo = str(row_marca['Resumo'])

            group_string += " " + resumo_limpo
            # ============================================

            document.add_paragraph(group_string)
            document.add_paragraph("")

        # --- Seção para resumos de Marca - NOTÍCIAS IRRELEVANTES ---
        document.add_paragraph("")
        document.add_paragraph("--- CITAÇÕES ---")
        document.add_paragraph("")
        
        print(f"Processando {len(df_resumo_marca_irrelevantes)} resumos de Marca - CITAÇÕES...")
        for index, row_marca in df_resumo_marca_irrelevantes.iterrows():
            group_string = ""

            if 'Ids' not in row_marca or pd.isna(row_marca['Ids']):
                print(f"Aviso: Linha {index} no df_resumo_marca_irrelevantes não tem IDs válidos. Pulando.")
                continue

            for news_id_str in str(row_marca['Ids']).split(','):
                try:
                    news_id = int(news_id_str.strip())
                except ValueError:
                    print(f"Aviso: Não foi possível converter ID '{news_id_str}' para inteiro na linha {index} de df_resumo_marca. Pulando.")
                    continue

                news_info_marca = final_df_marca[final_df_marca['Id'] == news_id]
                if news_info_marca.empty:
                    print(f"Aviso: ID {news_id} não encontrado em final_df_marca para resumo de Marca. Pulando.")
                    continue

                news_info_marca = news_info_marca.iloc[0]
                w_veiculo_marca = news_info_marca['Veiculo'].title()
                w_url_marca = news_info_marca['UrlVisualizacao']

                short_url_marca = encurtar_url_seguro(w_url_marca, max_tentativas=3, delay=1)

                if news_id in final_df_small_marca_irrelevantes['Id'].values:
                    final_df_small_marca_irrelevantes.loc[final_df_small_marca_irrelevantes['Id'] == news_id, 'ShortURL'] = short_url_marca

                special_type = ""
                special_info = final_df_SPECIALS_small[final_df_SPECIALS_small['Id'] == news_id]

                if not special_info.empty:
                    canais_special = special_info.iloc[0]['Canais']
                    if isinstance(canais_special, list):
                        canais_str = ', '.join(map(str, canais_special))
                    else:
                        canais_str = str(canais_special)

                    if "Editoriais" in canais_str:
                        special_type = "Editorial"
                    elif "Colunistas" in canais_str:
                        special_type = "Colunista"
                    elif "1ª Página" in canais_str:
                        special_type = "Capa"

                if special_type:
                    group_string += f"{w_veiculo_marca} ({special_type} - {short_url_marca}), "
                else:
                    group_string += f"{w_veiculo_marca} ({short_url_marca}), "

            # ========== CORREÇÃO APLICADA AQUI TAMBÉM ==========
            # Limpar vírgula final e adicionar espaço (SEM \n)
            group_string = group_string.rstrip(', ')

            # Adicionar resumo na mesma linha
            if 'Resumo' not in row_marca or pd.isna(row_marca['Resumo']):
                print(f"Aviso: Linha {index} no df_resumo_marca_irrelevantes não tem Resumo. Adicionando placeholder.")
                resumo_limpo = "[Resumo não disponível]"
            else:
                resumo_limpo = str(row_marca['Resumo'])

            group_string += " " + resumo_limpo
            # ===================================================

            document.add_paragraph(group_string)
            document.add_paragraph("")

        # --- Seção de links das notícias por Marcas ---
        final_df_small_marca_combined = pd.concat([final_df_small_marca, final_df_small_marca_irrelevantes], ignore_index=True)

        print(f"\nVerificando final_df_small_marca_combined antes da limpeza:")
        print(f"  Tem {len(final_df_small_marca_combined)} linhas.")
        
        if not final_df_small_marca_combined.empty:
            valores_nulos = final_df_small_marca_combined['Canais'].isna().sum()
            print(f"  Valores nulos em 'Canais': {valores_nulos}")
            
            final_df_small_marca_combined = final_df_small_marca_combined.dropna(subset=['Canais'])
            final_df_small_marca_combined = final_df_small_marca_combined[
                final_df_small_marca_combined['Canais'].astype(str).str.strip() != ''
            ]
            final_df_small_marca_combined = final_df_small_marca_combined[
                final_df_small_marca_combined['Canais'].astype(str).str.strip() != 'nan'
            ]
            
            print(f"  Após limpeza: {len(final_df_small_marca_combined)} linhas.")

        print(f"Tem coluna 'Canais'? {'Canais' in final_df_small_marca_combined.columns}")

        if not final_df_small_marca_combined.empty and 'Canais' in final_df_small_marca_combined.columns:
            try:
                marcas_unicas = final_df_small_marca_combined['Canais'].unique()
                marcas_unicas = [marca for marca in marcas_unicas if str(marca).strip() != '' and str(marca).strip() != 'nan']
                
                order = [marca1, marca2] + [marca for marca in marcas_unicas if marca not in (marca1, marca2)]
                
                print(f"Ordem de classificação: {order}")
                
                if len(order) > 0:
                    final_df_small_marca_sorted = final_df_small_marca_combined.sort_values(
                        by=['Canais'], 
                        key=lambda x: pd.Categorical(x, categories=order, ordered=True)
                    )
                    print("Ordenação personalizada dos links por Marca aplicada.")
                else:
                    print("Nenhuma categoria válida encontrada. Usando ordenação padrão.")
                    final_df_small_marca_sorted = final_df_small_marca_combined.sort_values(by=['Canais'])
                    
            except Exception as e:
                print(f"Aviso: Falha na ordenação personalizada dos links por Marca ({type(e).__name__}: {e}).")
                print("Ordenando por Canais padrão.")
                final_df_small_marca_sorted = final_df_small_marca_combined.sort_values(by=['Canais'])

            document.add_paragraph("")
            document.add_paragraph("--- LINKS DAS NOTÍCIAS DE MARCA ---")
            document.add_paragraph("")

            current_marca = None

            for index, row_small_marca in final_df_small_marca_sorted.iterrows():
                marca = row_small_marca['Canais']
                if marca != current_marca:
                    if current_marca is not None:
                        document.add_paragraph("")
                    document.add_paragraph(f"*{marca}*")
                    current_marca = marca

                news_id_small = row_small_marca['Id']
                original_row_marca = final_df_marca[final_df_marca['Id'] == news_id_small]
                if original_row_marca.empty:
                    print(f"Aviso: ID {news_id_small} não encontrado em final_df_marca para links de Marca. Pulando este link.")
                    continue

                original_row_marca = original_row_marca.iloc[0]
                veiculo = original_row_marca['Veiculo'].title()
                titulo = original_row_marca['Titulo']
                canais_commodities = original_row_marca.get('CanaisCommodities', '')

                document.add_paragraph(f"{veiculo}: {titulo}")

                short_url = row_small_marca.get('ShortURL')
                if pd.isna(short_url) or not short_url:
                    short_url = original_row_marca.get('UrlVisualizacao', 'URL Não Encontrada')

                prefix = "Coluna - " if "Colunistas" in str(canais_commodities) else ""
                document.add_paragraph(f"{prefix}{short_url}")
                document.add_paragraph("*")
        else:
            print("DataFrame 'final_df_small_marca_combined' não encontrado, vazio ou sem a coluna 'Canais'. Pulando a seção de links por Marcas.")

    # 2. SEÇÃO DE SETOR (CONDICIONAL)
    if executar_setor:
        print("\n=== SEÇÃO 2: SETOR ===")
        
        print(f"\nVerificando df_resumo_setor:")
        print(f"  Tem {len(df_resumo_setor)} linhas.")
        print(f"  Está vazio? {df_resumo_setor.empty}")
        
        if not df_resumo_setor.empty:
            document.add_paragraph("")
            document.add_paragraph("--- NOTÍCIAS DE SETOR ---")
            
            if codigo_veiculo:
                nome_veiculo = obter_nome_veiculo(codigo_veiculo)
                document.add_paragraph(nome_veiculo)
                print(f"Adicionado título específico do veículo: {nome_veiculo}")
            
            document.add_paragraph("")

            # Agrupar por Tema para evitar blocos repetidos do mesmo Tema
            # Limpar strings de Tema e garantir que valores nulos sejam pulados
            df_resumo_setor = df_resumo_setor.copy()
            if 'Tema' in df_resumo_setor.columns:
                df_resumo_setor['Tema'] = df_resumo_setor['Tema'].astype(str).str.strip()
                df_resumo_setor = df_resumo_setor[df_resumo_setor['Tema'].notna() & (df_resumo_setor['Tema'] != 'nan') & (df_resumo_setor['Tema'] != '')]

            # Opcional: ordenar dentro de cada tema por RelevanceScore descendente se a coluna existir
            sort_within = False
            if 'RelevanceScore' in df_resumo_setor.columns:
                try:
                    df_resumo_setor['RelevanceScore'] = pd.to_numeric(df_resumo_setor['RelevanceScore'], errors='coerce')
                    sort_within = True
                except Exception:
                    sort_within = False

            # Iterar por grupo de Tema — garante que cada tema gera apenas um cabeçalho
            for tema, grupo in df_resumo_setor.groupby('Tema', sort=False):
                if grupo.empty:
                    continue

                document.add_paragraph(f"*{tema}*")

                if sort_within:
                    grupo = grupo.sort_values(by='RelevanceScore', ascending=False)

                for index, row_setor in grupo.iterrows():
                    if 'Id' not in row_setor or pd.isna(row_setor['Id']):
                        print(f"Aviso: Linha {index} no df_resumo_setor não tem Id válido. Pulando.")
                        continue

                    try:
                        news_id = int(str(row_setor['Id']).strip())
                    except ValueError:
                        print(f"Aviso: Não foi possível converter ID '{row_setor['Id']}' para inteiro na linha {index} de df_resumo_setor. Pulando.")
                        continue

                    news_info_setor = final_df_setor[final_df_setor['Id'] == news_id]
                    if news_info_setor.empty:
                        print(f"Aviso: ID {news_id} não encontrado em final_df_setor para resumo de Setor. Pulando.")
                        continue
                    news_info_setor = news_info_setor.iloc[0]

                    w_veiculo_setor = news_info_setor['Veiculo'].title()
                    w_titulo_setor = news_info_setor['Titulo']
                    w_url_setor = news_info_setor['UrlVisualizacao']

                    document.add_paragraph(f"{w_veiculo_setor}: {w_titulo_setor}")

                    if 'Resumo' not in row_setor or pd.isna(row_setor['Resumo']):
                        print(f"Aviso: Linha {index} no df_resumo_setor (Tema {tema}) não tem Resumo. Adicionando placeholder.")
                        document.add_paragraph("[Resumo não disponível]")
                    else:
                        resumo_bruto = str(row_setor['Resumo'])
                        document.add_paragraph(resumo_bruto)

                    short_url_setor = encurtar_url_seguro(w_url_setor, max_tentativas=3, delay=1)
                    document.add_paragraph(short_url_setor)

                    document.add_paragraph("*")

    # 3. SEÇÃO DE EDITORIAIS (CONDICIONAL)
    if executar_editoriais:
        print("\n=== SEÇÃO 3: EDITORIAIS ===")
        
        document.add_paragraph("")
        document.add_paragraph("--- EDITORIAIS ---")    
        document.add_paragraph("")

        processar_editoriais_integrado(final_df_editorial, document, opcao_selecionada)

    # 4. SEÇÃO DE FINALIZAÇÃO (SEMPRE EXECUTADA)
    print("\n=== SEÇÃO 4: FINALIZAÇÃO ===")
    
    print("\nProcessando documento antes da gravação para remover padrões...")

    paragraphs_to_remove_indices = []

    pattern_line_start_paren_end_palavras = r"^\s*\*\s*\(\s*\d+\s*palavras\s*\)\s*$"
    pattern_line_start_resumo = r"^\s*\*Resumo.*$"
    # FIX: Mais específico - captura apenas (NN palavras) onde NN são dígitos, evitando capturar datas
    pattern_parenthesized_palavras_general = r"\s*[\*\s]*\(\s*\d+\s*palavras\s*\)[\s\:\*]*"
    pattern_specific_resumo_prefixes = r"^\s*[\*\s]*Resumo\s*[:\*\s]*"

    for i, paragraph in enumerate(document.paragraphs):
        texto = paragraph.text

        if re.fullmatch(pattern_line_start_paren_end_palavras, texto, re.IGNORECASE):
            paragraphs_to_remove_indices.append(i)
            print(f"  Removendo linha {i} por corresponder ao padrão '* (... palavras)'")
            continue

        if re.fullmatch(pattern_line_start_resumo, texto, re.IGNORECASE):
            paragraphs_to_remove_indices.append(i)
            print(f"  Removendo linha {i} por começar com '*Resumo'")
            continue

        new_text = re.sub(pattern_parenthesized_palavras_general, '', texto, flags=re.IGNORECASE).strip()
        new_text = re.sub(pattern_specific_resumo_prefixes, '', new_text, flags=re.IGNORECASE).strip()

        if new_text != texto.strip():
            paragraph.text = new_text
            print(f"  Removido padrão '(... palavras)' ou prefixos de resumo na linha {i}")

        if "**" in paragraph.text:
            paragraph.text = paragraph.text.replace("**", "*")

    for i in sorted(paragraphs_to_remove_indices, reverse=True):
        p = document.paragraphs[i]._element
        p.getparent().remove(p)

    print(f"\nTotal de parágrafos removidos: {len(paragraphs_to_remove_indices)}")

    # ═══════════════════════════════════════════════════════════════
    # VALIDAÇÃO ANTES DE SALVAR
    # ═══════════════════════════════════════════════════════════════
    try:
        print("\n🔍 Validando integridade do documento antes de salvar...")
        # Tentar acessar todos os parágrafos para detectar corrupção
        for i, p in enumerate(document.paragraphs):
            _ = p.text  # Acessar o texto para validar
        print(f"✅ Documento validado: {len(document.paragraphs)} parágrafos íntegros")
    except Exception as e:
        print(f"❌ ERRO ao validar documento: {e}")
        print("⚠️ Tentando recuperar salvando sem as alterações...")
        # Se houver erro, tentar salvar mesmo assim
    
    # Salvar com tratamento de erro
    try:
        document.save(arq_resumo_final)
        print(f"Arquivo DOCX salvo em: {arq_resumo_final}")
        
        # Verificar se o arquivo foi criado e tem tamanho
        if os.path.exists(arq_resumo_final):
            tamanho = os.path.getsize(arq_resumo_final)
            print(f"✅ Arquivo criado com sucesso ({tamanho} bytes)")
        else:
            print(f"❌ ERRO: Arquivo não foi criado!")
    except Exception as e:
        print(f"❌ ERRO ao salvar documento: {e}")
        print(f"   Traceback completo: {traceback.format_exc()}")
        raise
    
    print(f"\n=== RELATÓRIO GERADO COM SUCESSO ===")
    print(f"Opção processada: {opcao_selecionada}")
    if codigo_veiculo:
        print(f"Veículo: {obter_nome_veiculo(codigo_veiculo)}")
    print("=" * 50)