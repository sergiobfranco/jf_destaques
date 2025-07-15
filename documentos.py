from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
import re

def configurar_estilo(document):
    estilo = document.styles['Normal']
    fonte = estilo.font
    fonte.name = 'Calibri'
    fonte.size = Pt(11)
    par_format = estilo.paragraph_format
    par_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    par_format.space_after = Pt(0)

def adicionar_hyperlink(paragraph, url, texto_exibido):
    part = paragraph.add_run(texto_exibido)
    part.font.color.rgb = RGBColor(0, 0, 255)
    part.font.underline = True
    return paragraph

def limpar_texto(texto):
    texto = re.sub(r"\*\(\d+ palavras.*?\)\*", "", texto)
    texto = re.sub(r"\*Resumo:?\*?", "", texto, flags=re.IGNORECASE)
    texto = re.sub(r"\*\*+", "*", texto)
    return texto.strip()

def gerar_documento(df, caminho_saida):
    document = Document()
    configurar_estilo(document)

    for _, row in df.iterrows():
        titulo = row.get("Titulo", "").strip()
        conteudo = row.get("Conteudo", "").strip()
        resumo = row.get("Resumo", "").strip() if "Resumo" in row else ""
        veiculo = row.get("Veiculo", "")
        cidade = row.get("Cidade", "")
        url = row.get("URL", "")

        document.add_paragraph(f"{veiculo}/{cidade}: {titulo}")
        if resumo:
            document.add_paragraph(limpar_texto(resumo))
        else:
            document.add_paragraph(limpar_texto(conteudo))

        if url:
            adicionar_hyperlink(document.add_paragraph(), url, url)

        document.add_paragraph("*")  # separador visual

    document.save(caminho_saida)
    print(f"📄 Documento salvo em: {caminho_saida}")
