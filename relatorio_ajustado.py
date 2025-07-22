# Instalar as bibliotecas necessárias no Google Colab
#!pip install python-docx

import re
import os
from docx import Document
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement, qn
#from config import arq_resumo_final_ajustado, arq_resumo_final
from datetime import datetime

def adicionar_hyperlink(paragraph, url, texto_display):
    """
    Adiciona um hyperlink a um parágrafo no documento Word
    """
    # Criar o elemento hyperlink
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), paragraph.part.relate_to(url,
                                                       "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                                                       is_external=True))

    # Criar o run com o texto do link
    new_run = OxmlElement('w:r')

    # Configurar propriedades do texto (cor azul, sublinhado)
    rPr = OxmlElement('w:rPr')

    # Cor azul
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    # Sublinhado
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    new_run.append(rPr)

    # Adicionar o texto
    new_run.text = texto_display
    hyperlink.append(new_run)

    return hyperlink

def processar_urls_em_paragrafo(paragraph):
    """
    Processa um parágrafo, convertendo URLs em hyperlinks
    """
    texto_completo = paragraph.text

    # REGEX CORRIGIDA - Captura URLs completas incluindo parâmetros, fragmentos, etc.
    # Esta regex é mais robusta e captura URLs até encontrar espaço ou fim de linha
    padrao_url = r'https?://[^\s<>"{}|\\^`\[\]]+[^\s<>"{}|\\^`\[\].,;:!?)]'

    # Alternativa ainda mais simples e eficaz:
    # padrao_url = r'https?://\S+'

    urls_encontradas = re.findall(padrao_url, texto_completo)

    # Remover URLs duplicadas mantendo a ordem
    urls_unicas = []
    for url in urls_encontradas:
        if url not in urls_unicas:
            urls_unicas.append(url)

    if urls_unicas:
        print(f"   🔗 Encontradas {len(urls_unicas)} URLs: {urls_unicas[:2]}{'...' if len(urls_unicas) > 2 else ''}")

        # Limpar o parágrafo atual
        paragraph.clear()

        # Dividir o texto pelas URLs
        texto_restante = texto_completo

        for url in urls_unicas:
            # Encontrar a posição da URL no texto
            if url in texto_restante:
                partes = texto_restante.split(url, 1)

                if len(partes) == 2:
                    # Adicionar texto antes da URL
                    if partes[0]:
                        paragraph.add_run(partes[0])

                    # Adicionar hyperlink
                    hyperlink_element = adicionar_hyperlink(paragraph, url, url)
                    paragraph._p.append(hyperlink_element)

                    # Continuar com o resto do texto
                    texto_restante = partes[1]

        # Adicionar texto restante após a última URL
        if texto_restante:
            paragraph.add_run(texto_restante)

        return True

    return False

def converter_urls_docx_para_hyperlinks(arquivo_entrada, arquivo_saida):
    """
    Converte URLs em hyperlinks em um arquivo DOCX
    """
    try:
        # Abrir o documento
        print(f"📖 Abrindo arquivo: {arquivo_entrada}")
        doc = Document(arquivo_entrada)

        urls_convertidas = 0
        paragrafos_processados = 0
        paragrafos_com_urls = 0

        # Processar todos os parágrafos
        print("🔄 Processando parágrafos...")
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # Ignorar parágrafos vazios
                if processar_urls_em_paragrafo(paragraph):
                    paragrafos_com_urls += 1
                paragrafos_processados += 1

                # Mostrar progresso a cada 25 parágrafos
                if paragrafos_processados % 25 == 0:
                    print(f"   📄 Processados {paragrafos_processados} parágrafos ({paragrafos_com_urls} com URLs)...")

        # Processar tabelas (se houver)
        print("📊 Processando tabelas...")
        tabelas_processadas = 0
        celulas_com_urls = 0

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            if processar_urls_em_paragrafo(paragraph):
                                celulas_com_urls += 1
            tabelas_processadas += 1

        # Salvar o documento processado
        print(f"💾 Salvando arquivo: {arquivo_saida}")
        doc.save(arquivo_saida)

        print(f"\n✅ Conversão concluída com sucesso!")
        print(f"📊 Estatísticas:")
        print(f"   - Parágrafos processados: {paragrafos_processados}")
        print(f"   - Parágrafos com URLs: {paragrafos_com_urls}")
        print(f"   - Tabelas processadas: {tabelas_processadas}")
        print(f"   - Células com URLs: {celulas_com_urls}")
        print(f"   - Arquivo salvo como: {arquivo_saida}")

        return True

    except FileNotFoundError:
        print(f"❌ Erro: Arquivo '{arquivo_entrada}' não encontrado!")
        print("Verifique se o arquivo existe no diretório atual.")
        return False

    except Exception as e:
        print(f"❌ Erro durante o processamento: {str(e)}")
        print(f"Detalhes do erro: {type(e).__name__}")
        return False

def metodo_alternativo_melhorado(arquivo_entrada, arquivo_saida):
    """
    Método alternativo melhorado - substitui URLs por texto com formatação
    """
    try:
        print(f"📖 Método alternativo melhorado - Abrindo arquivo: {arquivo_entrada}")
        doc = Document(arquivo_entrada)

        total_urls_encontradas = 0
        paragrafos_processados = 0

        # Regex melhorada para capturar URLs completas
        padrao_url = r'https?://[^\s<>"{}|\\^`\[\]]+[^\s<>"{}|\\^`\[\].,;:!?)]'

        # Processar parágrafos
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue

            texto_original = paragraph.text
            urls_no_texto = re.findall(padrao_url, texto_original)

            if urls_no_texto:
                print(f"   🔗 Parágrafo {paragrafos_processados + 1}: {len(urls_no_texto)} URLs encontradas")
                total_urls_encontradas += len(urls_no_texto)

                # Limpar o parágrafo
                paragraph.clear()

                # Reconstruir o parágrafo com formatação
                texto_restante = texto_original

                for url in urls_no_texto:
                    if url in texto_restante:
                        # Dividir o texto pela URL
                        partes = texto_restante.split(url, 1)
                        if len(partes) == 2:
                            # Adicionar texto antes da URL
                            if partes[0]:
                                paragraph.add_run(partes[0])

                            # Adicionar URL com formatação especial
                            run_url = paragraph.add_run(url)
                            run_url.font.color.rgb = RGBColor(0, 0, 255)  # Azul
                            run_url.underline = True

                            # Continuar com o resto
                            texto_restante = partes[1]

                # Adicionar texto restante
                if texto_restante:
                    paragraph.add_run(texto_restante)

            paragrafos_processados += 1

        # Processar tabelas também
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if not paragraph.text.strip():
                            continue

                        texto_original = paragraph.text
                        urls_no_texto = re.findall(padrao_url, texto_original)

                        if urls_no_texto:
                            total_urls_encontradas += len(urls_no_texto)
                            paragraph.clear()

                            texto_restante = texto_original
                            for url in urls_no_texto:
                                if url in texto_restante:
                                    partes = texto_restante.split(url, 1)
                                    if len(partes) == 2:
                                        if partes[0]:
                                            paragraph.add_run(partes[0])

                                        run_url = paragraph.add_run(url)
                                        run_url.font.color.rgb = RGBColor(0, 0, 255)
                                        run_url.underline = True

                                        texto_restante = partes[1]

                            if texto_restante:
                                paragraph.add_run(texto_restante)

        # Salvar documento
        doc.save(arquivo_saida)

        print(f"\n✅ Método alternativo concluído!")
        print(f"📊 Estatísticas:")
        print(f"   - Parágrafos processados: {paragrafos_processados}")
        print(f"   - Total de URLs formatadas: {total_urls_encontradas}")
        print(f"💾 Arquivo salvo como: {arquivo_saida}")

        return True

    except Exception as e:
        print(f"❌ Erro no método alternativo: {str(e)}")
        return False

def testar_regex():
    """
    Função para testar a regex com URLs de exemplo
    """
    print("🧪 Testando regex com URLs de exemplo...")

    # URLs de teste
    urls_teste = [
        "https://tinyurl.com/2aymnjlf",
        "https://www.google.com/search?q=python",
        "http://example.com/path/to/file.html",
        "https://github.com/user/repo#readme",
        "https://site.com/page?param1=value1&param2=value2"
    ]

    # Regex melhorada
    padrao_url = r'https?://[^\s<>"{}|\\^`\[\]]+[^\s<>"{}|\\^`\[\].,;:!?)]'

    for url in urls_teste:
        match = re.search(padrao_url, url)
        if match:
            print(f"   ✅ {url} -> Capturado: {match.group()}")
        else:
            print(f"   ❌ {url} -> Não capturado")

    print("\n" + "="*50)

def gerar_versao_ajustada(arquivo_entrada):
    """
    Função principal para executar a conversão
    """
    # Nomes dos arquivos (ajuste conforme necessário)
    #arquivo_entrada = arq_resumo_final  # Seu arquivo original

    # Gerar novo nome de arquivo com timestamp antes do sufixo .docx
    base, ext = os.path.splitext(arquivo_entrada)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    arq_resumo_final_ajustado = f"{base}_{timestamp}{ext}"
 
    arquivo_saida = arq_resumo_final_ajustado  # Arquivo com hyperlinks

    print("🚀 Iniciando conversão de URLs para hyperlinks...")
    print("=" * 60)

    # Testar regex primeiro
    testar_regex()

    # Tentar primeiro método (hyperlinks verdadeiros)
    print("\n🔗 Tentando método com hyperlinks verdadeiros...")
    sucesso = converter_urls_docx_para_hyperlinks(arquivo_entrada, arquivo_saida)

    if not sucesso:
        print("\n🔄 Tentando método alternativo melhorado...")
        print("-" * 50)
        arquivo_saida_alt = "arq_resumo_final_formatado.docx"
        sucesso = metodo_alternativo_melhorado(arquivo_entrada, arquivo_saida_alt)

    if sucesso:
        print("\n🎉 Processo finalizado com sucesso!")
        print("\n💡 Dicas:")
        print("   - Abra o arquivo no Word para verificar os hyperlinks")
        print("   - Os links devem estar em azul e sublinhados")
        print("   - Teste alguns links clicando neles")
        print("   - Se algum link ainda estiver truncado, execute novamente")
    else:
        print("\n❌ Não foi possível completar a conversão.")
        print("Verifique se o arquivo de entrada existe e não está corrompido.")

